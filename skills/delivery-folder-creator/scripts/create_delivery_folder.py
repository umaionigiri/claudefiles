"""Create the monthly delivery folder for 大阪市統合PF.

Operates step-by-step so the calling agent (Claude) can ask the user
"proceed to next step?" between operations (smoke-test style).

Usage:
    python create_delivery_folder.py compute   --month 2026-05 --base "<base>"
    python create_delivery_folder.py step1     --month 2026-05 --base "<base>"  [--n 8] [--today YYYYMMDD]
    python create_delivery_folder.py step2     ...   # subfolder hierarchy
    python create_delivery_folder.py step3     ...   # copy 会議資料 (excludes 議事録)
    python create_delivery_folder.py step4     ...   # extract 議事録
    python create_delivery_folder.py step5     ...   # copy リスク管理簿 / 課題管理簿
    python create_delivery_folder.py step6     ...   # copy 稼働状況報告書 (YYYY年M月度)
    python create_delivery_folder.py step7     ...   # 検査仕様書
    python create_delivery_folder.py step8     ...   # 業務完了報告
    python create_delivery_folder.py step9     ...   # 統合PF請求書

Each step prints a JSON line with the result on stdout. `compute` returns the
parameters that subsequent steps need; pass them back via --n / --today so the
script stays stateless.
"""
from __future__ import annotations

import argparse
import calendar
import datetime as dt
import json
import re
import shutil
import sys
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

import docx

SCRIPT_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = SCRIPT_DIR.parent / "assets" / "templates"

KENSA_TPL = TEMPLATE_DIR / "kensa_yoshiki.xlsx"
GYOMU_TPL = TEMPLATE_DIR / "gyomu_kanryou.docx"
SEIKYUU_TPL = TEMPLATE_DIR / "seikyuu.docx"

# Customer-managed sources for meeting folders. READ-ONLY (must never be modified).
# Each entry maps a meeting type to:
#   src    — the customer-side folder containing 第N回<label>_YYYYMMDD subfolders
#   re     — regex matching folder names; group(1)=N, group(2)=YYYYMMDD
#   label  — name used both for the source folder convention and the destination
#            sub-bucket under 会議資料/ and 議事録/ in the delivery tree
CUSTOMER_DOC_ROOT = Path(
    "/mnt/c/Users/shoken.ohshiro/osakacitycommunication/"
    "B2B_大阪市統合プラットフォーム導入・構築 - ドキュメント"
)
# Separate customer SharePoint area for cross-service liaison documents.
# READ-ONLY (must never be modified — see references/delivery-rules.md).
CUSTOMER_RENRAKU_ROOT = Path(
    "/mnt/c/Users/shoken.ohshiro/osakacitycommunication/"
    "B2B_大阪市統合プラットフォーム連絡チーム - ドキュメント"
)
MEETING_SOURCES: list[dict] = [
    {
        "label": "作業部会",
        "src": CUSTOMER_DOC_ROOT / "01_作業部会",
        "re": re.compile(r"^第(\d+)回作業部会_(\d{8})"),
    },
    {
        "label": "進捗会議",
        "src": CUSTOMER_DOC_ROOT / "02_進捗会議",
        "re": re.compile(r"^第(\d+)回進捗会議_(\d{8})"),
    },
]

# Token used to detect 議事録 files inside meeting folders (substring match).
MINUTES_KEYWORD = "議事録"
# 1.0 release marker. Files whose name contains this token are considered final.
MINUTES_FINAL_TAG = "v1.00"

# Management ledgers copied from customer environments (read-only) into
# specific destination subfolders. Each entry:
#   src      — absolute source path (xlsx)
#   dest_sub — destination subfolder relative to 第n回中間分成果物/
#              (e.g. "プロジェクト管理/リスク管理表" or
#               "本稼働前対応/事業者間質問管理表")
#   stem     — filename stem; the destination becomes "<stem>_<YYYYMMDD>.xlsx"
#              where YYYYMMDD is the delivery-month end date.
MANAGEMENT_FILES: list[dict] = [
    {
        "src": CUSTOMER_DOC_ROOT / "General" / "00_共通" / "00_プロジェクト管理"
               / "30_リスク管理" / "【統合PF】全体_リスク管理簿.xlsx",
        "dest_sub": "プロジェクト管理/リスク管理表",
        "stem": "【統合PF】全体_リスク管理簿",
    },
    {
        "src": CUSTOMER_DOC_ROOT / "General" / "00_共通" / "00_プロジェクト管理"
               / "40_課題・ToDo管理" / "【統合PF】全体_課題・ToDo管理簿.xlsx",
        "dest_sub": "プロジェクト管理/課題管理表",
        "stem": "【統合PF】全体_課題・ToDo管理簿",
    },
    {
        # Source lives in the separate 連絡チーム SharePoint area (READ-ONLY).
        "src": CUSTOMER_RENRAKU_ROOT / "03_統合PFサービス横断会議連絡" / "00_マスタ"
               / "【共同編集】【統合PF】サービス横断会議_連絡台帳.xlsx",
        "dest_sub": "本稼働前対応/事業者間質問管理表",
        "stem": "【共同編集】【統合PF】サービス横断会議_連絡台帳",
    },
]

# 稼働状況報告書: monthly subfolders named "YYYY年M月度" (no zero-padding).
KASOU_SRC = CUSTOMER_DOC_ROOT / "General" / "02_ガバナンス検討" / "09_稼働状況報告書"
KASOU_RE = re.compile(r"^(\d{4})年(\d{1,2})月度")

ROOT_RE = re.compile(r"^\d{8}_(?:第)?(\d+)回目$")

SUBFOLDERS: list[str] = [
    "プロジェクト管理/リスク管理表",
    "プロジェクト管理/課題管理表",
    "プロジェクト管理/会議資料/作業部会",
    "プロジェクト管理/会議資料/進捗会議",
    "プロジェクト管理/議事録/作業部会",
    "プロジェクト管理/議事録/進捗会議",
    "ガバナンス検討",
    "本稼働前対応/事業者間質問管理表",
    "本稼働前対応/稼働状況報告書",
    "本稼働前対応/本稼働前対応計画書",
]


def reiwa_year(year: int, month: int) -> int:
    """Return 令和 fiscal-year number. 4月始まり: 4-12月→西暦-2018, 1-3月→西暦-2019."""
    return year - 2018 if month >= 4 else year - 2019


def reiwa_calendar_year(year: int) -> int:
    """Return 令和 calendar-year number for a calendar year (令和元年=2019)."""
    return year - 2018


def month_end(year: int, month: int) -> dt.date:
    last = calendar.monthrange(year, month)[1]
    return dt.date(year, month, last)


def next_month_15(year: int, month: int) -> dt.date:
    if month == 12:
        return dt.date(year + 1, 1, 15)
    return dt.date(year, month + 1, 15)


def detect_next_n(base: Path) -> int:
    if not base.exists():
        return 1
    nums = []
    for entry in base.iterdir():
        if not entry.is_dir():
            continue
        m = ROOT_RE.match(entry.name)
        if m:
            nums.append(int(m.group(1)))
    return (max(nums) + 1) if nums else 1


def parse_month(s: str) -> tuple[int, int]:
    """Accept '2026-05', '2026/5', '202605', '5月' (current year)."""
    s = s.strip()
    m = re.match(r"^(\d{4})[-/年]?(\d{1,2})月?$", s)
    if m:
        return int(m.group(1)), int(m.group(2))
    m = re.match(r"^(\d{4})(\d{2})$", s)
    if m:
        return int(m.group(1)), int(m.group(2))
    m = re.match(r"^(\d{1,2})月?$", s)
    if m:
        today = dt.date.today()
        return today.year, int(m.group(1))
    raise ValueError(f"Cannot parse month: {s!r}")


def compute_params(month: str, base: Path, today_str: str | None = None,
                   n_override: int | None = None) -> dict:
    year, mon = parse_month(month)
    today = dt.datetime.strptime(today_str, "%Y%m%d").date() if today_str else dt.date.today()
    n = n_override if n_override is not None else detect_next_n(base)
    me = month_end(year, mon)
    nm = next_month_15(year, mon)

    # Root folder is named after the *delivery-month end date*, not the
    # execution date — so April-delivery folders look like 20260430_第N回目
    # regardless of when the script actually runs.
    root_name = f"{me.strftime('%Y%m%d')}_第{n}回目"
    root_dir = base / root_name
    seika_dir = root_dir / f"第{n}回中間分成果物"

    return {
        "year": year,
        "month": mon,
        "n": n,
        "today_str": today.strftime("%Y%m%d"),
        "month_end_str": me.strftime("%Y%m%d"),
        "month_end_iso": me.isoformat(),
        "next_month15_iso": nm.isoformat(),
        "reiwa_fiscal": reiwa_year(year, mon),
        "reiwa_month_end": reiwa_calendar_year(me.year),
        "reiwa_next_month15": reiwa_calendar_year(nm.year),
        "base": str(base),
        "root_name": root_name,
        "root_dir": str(root_dir),
        "seika_dir": str(seika_dir),
        "kensa_filename": f"{me.strftime('%Y%m%d')}_検査仕様書（様式）.xlsx",
        "gyomu_filename": "業務完了報告.docx",
        "seikyuu_filename": f"統合PF請求書(第{n}回中間分業務).docx",
    }


# ---- step 1: root folder --------------------------------------------------
def step1_create_root(params: dict) -> dict:
    root = Path(params["root_dir"])
    if root.exists():
        return {"step": 1, "status": "exists", "root_dir": str(root)}
    root.mkdir(parents=True, exist_ok=False)
    return {"step": 1, "status": "created", "root_dir": str(root)}


# ---- step 2: subfolders ---------------------------------------------------
def step2_create_subfolders(params: dict) -> dict:
    seika = Path(params["seika_dir"])
    seika.mkdir(parents=True, exist_ok=True)
    created = []
    for rel in SUBFOLDERS:
        p = seika / rel
        p.mkdir(parents=True, exist_ok=True)
        created.append(str(p))
    return {"step": 2, "status": "ok", "seika_dir": str(seika), "folders_created": created}


# ---- step 3: 会議資料コピー / step 4: 議事録抽出 (作業部会 + 進捗会議) ------
def find_meetings_for_month(
    src: Path, name_re: re.Pattern, year: int, month: int
) -> list[Path]:
    """List meeting folders (作業部会 / 進捗会議) whose embedded date falls in
    the target year-month. Read-only: never mutates the source directory.
    """
    if not src.exists():
        return []
    target_prefix = f"{year:04d}{month:02d}"
    matches: list[tuple[int, Path]] = []
    for entry in src.iterdir():
        if not entry.is_dir():
            continue
        m = name_re.match(entry.name)
        if not m:
            continue
        if m.group(2).startswith(target_prefix):
            matches.append((int(m.group(1)), entry))
    matches.sort(key=lambda x: x[0])
    return [p for _, p in matches]


def _ignore_minutes(_dirpath: str, names: list[str]) -> list[str]:
    """shutil.copytree ignore callback: skip any entry whose name contains 議事録.

    Applies at every directory level (top + subfolders like old/), so that the
    会議資料 destination tree never contains a minutes file.
    """
    return [n for n in names if MINUTES_KEYWORD in n]


def _copy_meeting_files_for_label(
    params: dict, label: str, src_root: Path, name_re: re.Pattern
) -> dict:
    """Step 3 worker: copy meeting folders for one label (作業部会/進捗会議),
    excluding minutes files. Returns a per-label result dict.
    """
    dst_dir = Path(params["seika_dir"]) / "プロジェクト管理" / "会議資料" / label
    dst_dir.mkdir(parents=True, exist_ok=True)

    if not src_root.exists():
        return {"label": label, "status": "skipped",
                "reason": f"Source not found: {src_root}", "dst_dir": str(dst_dir)}

    meetings = find_meetings_for_month(src_root, name_re, params["year"], params["month"])
    if not meetings:
        return {"label": label, "status": "no_match",
                "src": str(src_root), "dst_dir": str(dst_dir)}

    copied: list[str] = []
    skipped: list[str] = []
    for src_meeting in meetings:
        dst = dst_dir / src_meeting.name
        if dst.exists():
            skipped.append(src_meeting.name)
            continue
        # copytree copies recursively WITHOUT touching the source. `ignore`
        # strips out 議事録 files (those are placed in 議事録/ by step4).
        shutil.copytree(src_meeting, dst, ignore=_ignore_minutes)
        copied.append(src_meeting.name)
    return {"label": label, "status": "ok", "dst_dir": str(dst_dir),
            "copied": copied, "skipped_already_exists": skipped}


def step3_copy_meetings(params: dict) -> dict:
    """Copy 作業部会 / 進捗会議 folders for the delivery month, with 議事録
    files excluded (議事録 is handled in step4).

    Source: customer environment (READ-ONLY).
    Dest:   <root>/第n回中間分成果物/プロジェクト管理/会議資料/{作業部会, 進捗会議}/
    """
    per_label = [
        _copy_meeting_files_for_label(params, s["label"], s["src"], s["re"])
        for s in MEETING_SOURCES
    ]
    return {"step": 3, "status": "ok", "minutes_excluded": True, "per_label": per_label}


PLACEHOLDER_NAME = "※確認中のため次月納品.txt"


def _find_previous_delivery_root(base: Path, current_n: int) -> Path | None:
    """Locate the previous delivery root folder (n-1) under base. Returns None if absent."""
    if current_n <= 1 or not base.exists():
        return None
    target_n = current_n - 1
    for entry in base.iterdir():
        if not entry.is_dir():
            continue
        m = ROOT_RE.match(entry.name)
        if m and int(m.group(1)) == target_n:
            return entry
    return None


def _find_pending_meetings(prev_root: Path, label: str) -> list[str]:
    """Return meeting folder names that should be reconsidered as carry-overs.

    Two cases are detected:
    A) Explicit placeholder: 議事録/<label>/<meeting>/ exists with
       ※確認中のため次月納品.txt — minutes were known-pending last time.
    B) Implicit miss: 会議資料/<label>/<meeting>/ existed in the previous
       delivery but 議事録/<label>/<meeting>/ never got created at all (e.g.
       v1.00 was not yet confirmed at the time of the previous delivery and an
       earlier version of this script did not place a placeholder for it).
       These also need to be carried over once the customer environment now
       holds a finalized v1.00.
    """
    pending: list[str] = []
    minutes_dir: Path | None = None
    materials_dir: Path | None = None
    for child in prev_root.iterdir():
        if child.is_dir() and child.name.endswith("中間分成果物"):
            cand_minutes = child / "プロジェクト管理" / "議事録" / label
            cand_materials = child / "プロジェクト管理" / "会議資料" / label
            if cand_minutes.exists():
                minutes_dir = cand_minutes
            if cand_materials.exists():
                materials_dir = cand_materials
            break

    # Case A — explicit placeholder.
    if minutes_dir is not None:
        for meeting_dir in minutes_dir.iterdir():
            if meeting_dir.is_dir() and (meeting_dir / PLACEHOLDER_NAME).exists():
                pending.append(meeting_dir.name)

    # Case B — present in 会議資料 but missing from 議事録.
    if materials_dir is not None:
        existing_minute_names = (
            {p.name for p in minutes_dir.iterdir() if p.is_dir()}
            if minutes_dir is not None else set()
        )
        for material_meeting in materials_dir.iterdir():
            if not material_meeting.is_dir():
                continue
            name = material_meeting.name
            if name in existing_minute_names:
                continue
            if name in pending:
                continue
            pending.append(name)

    return pending


def _copy_minutes_for_label(
    params: dict, label: str, src_root: Path, name_re: re.Pattern
) -> dict:
    """Step 4 worker: pick each meeting's v1.00 議事録 (or place a placeholder)
    under 議事録/<label>/, and additionally pull in any carry-overs from the
    previous delivery's pending placeholders that are now finalized (v1.00).
    """
    dst_root = Path(params["seika_dir"]) / "プロジェクト管理" / "議事録" / label
    dst_root.mkdir(parents=True, exist_ok=True)

    if not src_root.exists():
        return {"label": label, "status": "skipped",
                "reason": f"Source not found: {src_root}", "dst_dir": str(dst_root)}

    placeholder_body = (
        f"このフォルダには {MINUTES_FINAL_TAG} 版の議事録がまだ存在しないため、"
        "次月の納品で改めて格納します。\n"
    )

    def process_meeting(src_meeting: Path, action_when_copied: str) -> dict:
        """Common logic for both current-month meetings and carry-over meetings.
        Returns a result entry. `action_when_copied` is either "copied" (current
        month) or "carried_over" (pulled from previous delivery's placeholder)."""
        per_dst = dst_root / src_meeting.name

        top_minutes = [
            p for p in src_meeting.iterdir()
            if p.is_file() and MINUTES_KEYWORD in p.name
        ]
        final_minutes = [p for p in top_minutes if MINUTES_FINAL_TAG in p.name]

        if final_minutes:
            per_dst.mkdir(parents=True, exist_ok=True)
            entry = {"meeting": src_meeting.name, "action": action_when_copied, "files": []}
            for src_file in final_minutes:
                dst_file = per_dst / src_file.name
                if not dst_file.exists():
                    shutil.copyfile(src_file, dst_file)
                entry["files"].append(src_file.name)
            return entry

        if action_when_copied == "carried_over":
            # Carry-over still pending: don't add a fresh placeholder here —
            # the previous delivery already holds one. Just record the state.
            return {
                "meeting": src_meeting.name,
                "action": "carried_over_still_pending",
                "reason": f"previous delivery placeholder; v{MINUTES_FINAL_TAG} still not found",
                "non_final_seen": [p.name for p in top_minutes],
            }

        # Current-month meeting without v1.00 → place a placeholder.
        per_dst.mkdir(parents=True, exist_ok=True)
        placeholder_path = per_dst / PLACEHOLDER_NAME
        if not placeholder_path.exists():
            placeholder_path.write_text(placeholder_body, encoding="utf-8")
        return {
            "meeting": src_meeting.name, "action": "placeholder",
            "reason": ("no minutes file found"
                       if not top_minutes
                       else f"no file with '{MINUTES_FINAL_TAG}' tag"),
            "non_final_seen": [p.name for p in top_minutes],
            "placeholder": PLACEHOLDER_NAME,
        }

    results: list[dict] = []

    # 1) Current-month meetings.
    current = find_meetings_for_month(src_root, name_re, params["year"], params["month"])
    for src_meeting in current:
        results.append(process_meeting(src_meeting, "copied"))

    # 2) Carry-overs from previous delivery's placeholders.
    prev_root = _find_previous_delivery_root(Path(params["base"]), params["n"])
    current_names = {p.name for p in current}
    if prev_root is not None:
        for pending_name in _find_pending_meetings(prev_root, label):
            if pending_name in current_names:
                # Already handled as part of the current month — skip.
                continue
            src_meeting = src_root / pending_name
            if not src_meeting.exists():
                results.append({
                    "meeting": pending_name,
                    "action": "carried_over_source_missing",
                    "reason": "previous delivery placeholder; source folder no longer exists",
                })
                continue
            results.append(process_meeting(src_meeting, "carried_over"))

    return {
        "label": label, "status": "ok", "dst_dir": str(dst_root),
        "previous_delivery": str(prev_root) if prev_root else None,
        "results": results,
    }


def step4_copy_minutes(params: dict) -> dict:
    """Place v1.00 議事録 files (or placeholders) under 議事録/{作業部会, 進捗会議}/."""
    per_label = [
        _copy_minutes_for_label(params, s["label"], s["src"], s["re"])
        for s in MEETING_SOURCES
    ]
    return {"step": 4, "status": "ok", "per_label": per_label}


# ---- step 5: 管理簿 (リスク管理表 / 課題管理表) のコピー --------------------
def step5_copy_management_files(params: dict) -> dict:
    """Copy management ledgers from the customer environment to the delivery tree.

    For each entry in MANAGEMENT_FILES:
      src = customer-side xlsx (READ-ONLY, never mutated)
      dst = <root>/第n回中間分成果物/プロジェクト管理/<dest_sub>/<stem>_<YYYYMMDD>.xlsx
            where YYYYMMDD is the delivery-month end date.

    Pure file copy — no content edits, so xlsx structure is byte-for-byte
    preserved (no openpyxl rewrite, no Excel-recovery prompt).
    """
    suffix = params["month_end_str"]  # YYYYMMDD of delivery-month end
    seika = Path(params["seika_dir"])
    results: list[dict] = []

    for entry in MANAGEMENT_FILES:
        src: Path = entry["src"]
        # dest_sub is a relative path under 第n回中間分成果物/, e.g.
        # "プロジェクト管理/リスク管理表" or "本稼働前対応/事業者間質問管理表".
        dst_dir = seika / entry["dest_sub"]
        dst_dir.mkdir(parents=True, exist_ok=True)
        dst = dst_dir / f"{entry['stem']}_{suffix}.xlsx"

        if not src.exists():
            results.append({
                "stem": entry["stem"], "status": "missing",
                "src": str(src), "dst": str(dst),
            })
            continue
        if dst.exists():
            results.append({
                "stem": entry["stem"], "status": "exists", "dst": str(dst),
            })
            continue
        shutil.copyfile(src, dst)
        results.append({
            "stem": entry["stem"], "status": "copied",
            "src": str(src), "dst": str(dst),
        })

    return {"step": 5, "status": "ok", "results": results}


# ---- step 6: 稼働状況報告書コピー ------------------------------------------
def step6_copy_kasou_joukyou(params: dict) -> dict:
    """Copy the customer-managed 稼働状況報告書 folder for the **previous** month.

    Business rule: a delivery month covers the *previous* month's report —
    e.g. April delivery uses March's "<YYYY年3月度>" folder.

    Source: ガバナンス検討/09_稼働状況報告書/<YYYY年M月度>/  (READ-ONLY)
    Dest:   <root>/第n回中間分成果物/本稼働前対応/稼働状況報告書/<YYYY年M月度>/

    Folder names use no zero-padding (e.g. "2026年3月度", "2026年11月度").
    Multiple folders for the same month are all copied (rare but supported).
    """
    dst_dir = (
        Path(params["seika_dir"]) / "本稼働前対応" / "稼働状況報告書"
    )
    dst_dir.mkdir(parents=True, exist_ok=True)

    if not KASOU_SRC.exists():
        return {"step": 6, "status": "skipped",
                "reason": f"Source not found: {KASOU_SRC}", "dst_dir": str(dst_dir)}

    # Target = (delivery year, month) - 1 month.
    year = params["year"]
    month = params["month"]
    if month == 1:
        target_year, target_month = year - 1, 12
    else:
        target_year, target_month = year, month - 1

    matches: list[Path] = []
    for entry in KASOU_SRC.iterdir():
        if not entry.is_dir():
            continue
        m = KASOU_RE.match(entry.name)
        if m and int(m.group(1)) == target_year and int(m.group(2)) == target_month:
            matches.append(entry)
    matches.sort(key=lambda p: p.name)

    if not matches:
        return {"step": 6, "status": "no_match",
                "delivery_month": f"{year:04d}-{month:02d}",
                "target_month": f"{target_year:04d}-{target_month:02d}",
                "src": str(KASOU_SRC), "dst_dir": str(dst_dir)}

    copied: list[str] = []
    skipped: list[str] = []
    for src in matches:
        dst = dst_dir / src.name
        if dst.exists():
            skipped.append(src.name)
            continue
        # Pure copy. The customer source is never written.
        shutil.copytree(src, dst)
        copied.append(src.name)
    return {"step": 6, "status": "ok", "dst_dir": str(dst_dir),
            "delivery_month": f"{year:04d}-{month:02d}",
            "target_month": f"{target_year:04d}-{target_month:02d}",
            "copied": copied, "skipped_already_exists": skipped}


# ---- step 7: 検査仕様書 ---------------------------------------------------
SHEET_NAME_RE = re.compile(rb'(<sheet\b[^/>]*?\bname=")([^"]+)(")')


def rename_first_sheet_xlsx(path: Path, new_name: str) -> str:
    """Rename the first sheet inside an xlsx by editing only XML in the zip.

    Updates two things:
      1) <sheet name="OLD" .../> in xl/workbook.xml (the user-visible sheet name)
      2) Every "OLD!"/"'OLD'!" reference in xl/workbook.xml (definedNames) and
         in xl/worksheets/sheet*.xml (cell formulas), so internal references
         to the renamed sheet stay consistent. Without this Excel reports
         "ブックを回復しますか" because definedNames point at a non-existent sheet.

    Every other zip entry (printerSettings, comments, customXml, sharedStrings,
    threadedComments, drawings, _rels, etc.) is copied byte-for-byte. This avoids
    the data loss that occurs when openpyxl rewrites the workbook from its parsed
    model.
    """
    with zipfile.ZipFile(path, "r") as zin:
        wb_xml = zin.read("xl/workbook.xml")

    m = SHEET_NAME_RE.search(wb_xml)
    if not m:
        raise RuntimeError("Cannot find <sheet name=...> in workbook.xml")
    old_name = m.group(2).decode("utf-8")

    # 1) Replace the first <sheet name="..."> attribute.
    new_attr = m.group(1) + xml_escape(new_name).encode("utf-8") + m.group(3)
    new_wb_xml = wb_xml[: m.start()] + new_attr + wb_xml[m.end():]

    # 2) Replace internal sheet-name references. Excel writes them as either
    #    "Sheet!Cell" (no spaces) or "'Sheet Name'!Cell" (with spaces). Cover
    #    both. Apply the quoted form first so we don't double-replace.
    old_esc = xml_escape(old_name).encode("utf-8")
    new_esc = xml_escape(new_name).encode("utf-8")
    refs: list[tuple[bytes, bytes]] = [
        (b"'" + old_esc + b"'!", b"'" + new_esc + b"'!"),
        (old_esc + b"!", new_esc + b"!"),
    ]

    def patch(data: bytes) -> bytes:
        for old, new in refs:
            data = data.replace(old, new)
        return data

    new_wb_xml = patch(new_wb_xml)

    tmp_path = path.with_name(path.name + ".tmp")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp_path, "w") as zout:
        for info in zin.infolist():
            name = info.filename
            if name == "xl/workbook.xml":
                data = new_wb_xml
            elif name.startswith("xl/worksheets/") and name.endswith(".xml"):
                # Patch cell-formula references like ="OLD"!A1 in every sheet.
                data = patch(zin.read(name))
            else:
                data = zin.read(name)
            zout.writestr(info, data)
    shutil.move(tmp_path, path)
    return f"{old_name} -> {new_name}"


def step7_create_kensa(params: dict) -> dict:
    """Copy kensa template, rename first sheet to 第N回中間分 (XML-direct, no openpyxl)."""
    out = Path(params["root_dir"]) / params["kensa_filename"]
    if out.exists():
        return {"step": 7, "status": "exists", "path": str(out)}
    shutil.copyfile(KENSA_TPL, out)
    new_sheet_name = f"第{params['n']}回中間分"
    rename_info = rename_first_sheet_xlsx(out, new_sheet_name)
    return {"step": 7, "status": "created", "path": str(out), "sheet_renamed": rename_info}


# ---- step 8: 業務完了報告 -------------------------------------------------
def replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a paragraph's text while preserving the first run's formatting.

    All runs after the first are removed; the first run's `text` is set to
    `new_text`. This drops per-run inline formatting variations but preserves
    paragraph alignment, run font/size of the first run, and surrounding XML.
    """
    runs = paragraph.runs
    if not runs:
        # No runs: add a new run.
        paragraph.add_run(new_text)
        return
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    runs[0].text = new_text


def step8_create_gyomu(params: dict) -> dict:
    out = Path(params["root_dir"]) / params["gyomu_filename"]
    if out.exists():
        return {"step": 8, "status": "exists", "path": str(out)}
    shutil.copyfile(GYOMU_TPL, out)
    doc = docx.Document(out)

    me_iso = params["month_end_iso"]
    me_y, me_m, me_d = me_iso.split("-")
    reiwa_me = params["reiwa_month_end"]
    reiwa_fy = params["reiwa_fiscal"]
    n = params["n"]

    date_text = f"令和{reiwa_me}年{int(me_m)}月{int(me_d)}日　"
    contract_text = f"令和{reiwa_fy}年度　第{n}回中間分業務"

    # P1 = right-aligned date paragraph (per inspect_templates.py output)
    date_para = doc.paragraphs[1]
    if "令和" not in date_para.text:
        # Fallback: locate the first paragraph that contains 令和 and a 日.
        for p in doc.paragraphs:
            if "令和" in p.text and "日" in p.text:
                date_para = p
                break
    replace_paragraph_text(date_para, date_text)

    # Table 0 row 5 col 1 = 履行内容 cell (per inspect_templates.py output).
    table = doc.tables[0]
    contract_cell = table.cell(5, 1)
    if "履行内容" in table.cell(5, 0).text:
        # Replace the first paragraph of the cell.
        replace_paragraph_text(contract_cell.paragraphs[0], contract_text)

    doc.save(out)
    return {
        "step": 8, "status": "created", "path": str(out),
        "date_replaced": date_text,
        "contract_replaced": contract_text,
    }


# ---- step 9: 統合PF請求書 -------------------------------------------------
def step9_create_seikyuu(params: dict) -> dict:
    out = Path(params["root_dir"]) / params["seikyuu_filename"]
    if out.exists():
        return {"step": 9, "status": "exists", "path": str(out)}
    shutil.copyfile(SEIKYUU_TPL, out)
    doc = docx.Document(out)

    nm = dt.date.fromisoformat(params["next_month15_iso"])
    reiwa_nm = params["reiwa_next_month15"]
    date_text = f"　令和{reiwa_nm}年　{nm.month}月　{nm.day}日"

    # Table 0 row 2 (merged across 26 columns) = top-right invoice date.
    table = doc.tables[0]
    date_cell = table.cell(2, 0)
    replace_paragraph_text(date_cell.paragraphs[0], date_text)

    doc.save(out)
    return {"step": 9, "status": "created", "path": str(out), "date_replaced": date_text}


# ---- Summary md ------------------------------------------------------------
def _compact_meeting_range(names: list[str], num_re: re.Pattern) -> str:
    """Compress a list of meeting folder names into e.g. '第55-58回' or '第55,57回'."""
    nums: list[int] = []
    for n in names:
        m = num_re.match(n)
        if m:
            nums.append(int(m.group(1)))
    if not nums:
        return ""
    nums = sorted(set(nums))
    groups: list[tuple[int, int]] = []
    start = prev = nums[0]
    for x in nums[1:]:
        if x == prev + 1:
            prev = x
            continue
        groups.append((start, prev))
        start = prev = x
    groups.append((start, prev))
    parts = [f"{s}-{e}" if s != e else f"{s}" for s, e in groups]
    return f"第{','.join(parts)}回"


def _minutes_state_summary(label_result: dict, label: str) -> str:
    """Produce one-line summary of minutes status, e.g. '第55,56,57回 v1.00 + 第58回 ※確認中'."""
    finalized: list[str] = []
    pending: list[str] = []
    num_re = re.compile(rf"^第(\d+)回{label}_")
    for x in label_result.get("results", []) or []:
        action = x.get("action")
        # Both 'copied' (current month with v1.00) and 'carried_over' (previous
        # delivery's miss now finalized) mean a v1.00 file exists in the tree.
        if action in ("copied", "carried_over"):
            finalized.append(x["meeting"])
        else:
            # placeholder / carried_over_still_pending / carried_over_source_missing
            pending.append(x["meeting"])
    parts: list[str] = []
    if finalized:
        parts.append(f"{_compact_meeting_range(finalized, num_re)} v1.00")
    if pending:
        parts.append(f"{_compact_meeting_range(pending, num_re)} ※確認中")
    return " + ".join(parts) if parts else "(なし)"


def _build_tree_section(by_step: dict, p: dict) -> list[str]:
    """Return Markdown lines (inside a fenced ``` block) showing the delivery tree."""
    n = p.get("n", "")
    root_name = p.get("root_name", "")
    today_str = p.get("today_str", "")
    me_str = p.get("month_end_str", "")

    s7 = by_step.get(7, {})
    kensa_name = Path(s7.get("path", "")).name or f"{me_str}_検査仕様書（様式）.xlsx"

    s8 = by_step.get(8, {})
    gyomu_name = Path(s8.get("path", "")).name or "業務完了報告.docx"
    gyomu_note = ""
    if s8.get("date_replaced") and s8.get("contract_replaced"):
        gyomu_note = f" ({s8['date_replaced'].strip()} / {s8['contract_replaced']})"

    s9 = by_step.get(9, {})
    seikyuu_name = Path(s9.get("path", "")).name or f"統合PF請求書(第{n}回中間分業務).docx"
    seikyuu_note = f" ({s9.get('date_replaced', '').strip()})" if s9.get("date_replaced") else ""

    # Step 5 file names (by stem prefix → bucket)
    risk_name = kadai_name = renraku_name = ""
    for x in by_step.get(5, {}).get("results", []):
        dst_name = Path(x.get("dst", "")).name
        stem = x.get("stem", "")
        if "リスク管理簿" in stem:
            risk_name = dst_name
        elif "課題・ToDo管理簿" in stem:
            kadai_name = dst_name
        elif "サービス横断会議_連絡台帳" in stem:
            renraku_name = dst_name

    # Step 3 ranges per label — merge 'copied' (this run) with
    # 'skipped_already_exists' (idempotent re-run) so the tree always reflects
    # what's currently present in the delivery folder.
    s3 = {lab["label"]: lab for lab in by_step.get(3, {}).get("per_label", [])}
    sagyou_present = list(s3.get("作業部会", {}).get("copied", []) or []) + list(
        s3.get("作業部会", {}).get("skipped_already_exists", []) or []
    )
    shinchoku_present = list(s3.get("進捗会議", {}).get("copied", []) or []) + list(
        s3.get("進捗会議", {}).get("skipped_already_exists", []) or []
    )
    sagyou_range = _compact_meeting_range(
        sagyou_present, re.compile(r"^第(\d+)回作業部会_"),
    )
    shinchoku_range = _compact_meeting_range(
        shinchoku_present, re.compile(r"^第(\d+)回進捗会議_"),
    )

    # Step 4 minutes state
    s4 = {lab["label"]: lab for lab in by_step.get(4, {}).get("per_label", [])}
    minutes_sagyou = _minutes_state_summary(s4.get("作業部会", {}), "作業部会")
    minutes_shinchoku = _minutes_state_summary(s4.get("進捗会議", {}), "進捗会議")

    # Step 6 kasou folder names — also merge skipped_already_exists for re-runs.
    s6 = by_step.get(6, {})
    kasou_present = list(s6.get("copied", []) or []) + list(s6.get("skipped_already_exists", []) or [])
    kasou_str = ", ".join(kasou_present) if kasou_present else "(該当なし)"

    summary_md_name = f"納品作成サマリ_{today_str}.md"

    tree: list[str] = []
    tree.append("```")
    tree.append(f"{root_name}/")
    tree.append(f"├ {summary_md_name}            ← ★ このファイル")
    tree.append(f"├ {kensa_name}       (シート名「第{n}回中間分」)")
    tree.append(f"├ {gyomu_name}                     {gyomu_note}")
    tree.append(f"├ {seikyuu_name}     {seikyuu_note}")
    tree.append(f"└ 第{n}回中間分成果物/")
    tree.append(f"   ├ プロジェクト管理/")
    tree.append(f"   │   ├ リスク管理表/{risk_name}")
    tree.append(f"   │   ├ 課題管理表/{kadai_name}")
    tree.append(f"   │   ├ 会議資料/")
    tree.append(f"   │   │   ├ 作業部会/{{{sagyou_range}}}                     (議事録除外)")
    tree.append(f"   │   │   └ 進捗会議/{{{shinchoku_range}}}                     (議事録除外)")
    tree.append(f"   │   └ 議事録/")
    tree.append(f"   │       ├ 作業部会/{{{minutes_sagyou}}}")
    tree.append(f"   │       └ 進捗会議/{{{minutes_shinchoku}}}")
    tree.append(f"   ├ ガバナンス検討/                                  (空)")
    tree.append(f"   └ 本稼働前対応/")
    tree.append(f"       ├ 事業者間質問管理表/{renraku_name}")
    tree.append(f"       ├ 稼働状況報告書/{kasou_str}/                   (前月度コピー)")
    tree.append(f"       └ 本稼働前対応計画書/                           (空)")
    tree.append("```")
    return tree


def _label_meeting_summary(label_result: dict, label: str) -> tuple[int, list[str], list[dict]]:
    """For step3: return (copied_count, copied_meeting_names, []) — placeholders unused."""
    copied = label_result.get("copied", []) or []
    return len(copied), copied, []


def format_summary_md(params: dict, results: list[dict]) -> str:
    """Render a result list (from `all` or `summary` mode) into a Markdown
    summary doc.

    Reads naturally for non-technical readers — no jargon (copied/created/ok),
    uses ✅/⚠ icons and Japanese-only descriptions.

    `params` is the canonical compute-step output, used as a fallback for
    fields that are only populated on first-run (e.g. step 7-9 rename details
    are absent when those files already existed).
    """
    by_step = {r.get("step"): r for r in results}
    p = by_step.get(0, params or {})

    # Collect placeholders early so we can both render Step 4 and the
    # "次に対応してください" section consistently.
    placeholder_entries: list[tuple[str, str]] = []  # (label, meeting_name)
    for lab in by_step.get(4, {}).get("per_label", []):
        for x in lab.get("results", []) or []:
            if x.get("action") == "placeholder":
                placeholder_entries.append((lab.get("label", ""), x.get("meeting", "")))

    lines: list[str] = []
    lines.append("# 大阪市統合PF 月次納品 自動生成サマリ")
    lines.append("")
    lines.append("このファイルは納品フォルダ作成スキルが自動生成したサマリです。何がどこに作成・取り込まれたかを一目で確認できます。")
    lines.append("")
    lines.append("## 実行情報")
    lines.append("")
    lines.append("| 項目 | 値 |")
    lines.append("|-----|-----|")
    lines.append(f"| 実行日 | {p.get('today_str', '')} |")
    lines.append(f"| 納品月 | {p.get('year', '')}年{p.get('month', '')}月 |")
    lines.append(f"| 回数 | 第{p.get('n', '')}回目 |")
    lines.append(f"| 月末日（業務完了報告の日付） | {p.get('month_end_iso', '')} |")
    lines.append(f"| 翌月15日（請求書の日付） | {p.get('next_month15_iso', '')} |")
    lines.append(f"| 令和会計年度（履行内容） | 令和{p.get('reiwa_fiscal', '')}年度 |")
    lines.append("")
    lines.append("## 作成パス")
    lines.append(f"- 当月納品フォルダ: `{p.get('root_dir', '')}`")
    lines.append(f"- 中間成果物フォルダ: `{p.get('seika_dir', '')}`")
    lines.append("")
    lines.append("## 作成ツリー")
    lines.extend(_build_tree_section(by_step, p))
    lines.append("")
    lines.append("## 作業の内訳")
    lines.append("")
    lines.append("以下は、自動化が代行した9つの作業の内訳です。各項目の先頭アイコンは:")
    lines.append("")
    lines.append("- ✅ … 正常完了 / 📂 … フォルダ作成 / 📥 … 顧客環境からの取り込み / 📝 … 文書の自動生成 / ⚠ … 要確認")
    lines.append("")

    # ---- 1. ルートフォルダ
    s = by_step.get(1, {})
    lines.append("### 1. 当月納品フォルダの新規作成")
    if s.get("status") == "created":
        lines.append(f"📂 `{Path(s.get('root_dir', '')).name}` を新規作成しました。")
    elif s.get("status") == "exists":
        lines.append(f"📂 `{Path(s.get('root_dir', '')).name}` は既に存在していたため、上書きせず既存フォルダを使用しています。")
    lines.append("")

    # ---- 2. サブフォルダ階層
    s = by_step.get(2, {})
    folders_n = len(s.get("folders_created", []) or [])
    lines.append("### 2. サブフォルダ階層の作成")
    lines.append(f"📂 「プロジェクト管理 / 会議資料 / 議事録 / ガバナンス検討 / 本稼働前対応」など、合計 {folders_n} 個の空フォルダを作成しました。")
    lines.append("")

    # ---- 3. 会議資料の取り込み
    s3 = by_step.get(3, {})
    lines.append("### 3. 会議資料の取り込み（顧客環境 → 会議資料配下）")
    lines.append("お客さま環境（`B2B_導入・構築` / `01_作業部会`・`02_進捗会議`）の **読み取りのみ** で、納品月に該当する会議フォルダを丸ごとコピーしました。**議事録ファイルはこのフォルダには含めず**、次の Step 4 で別フォルダに配置しています。")
    lines.append("")
    for lab in s3.get("per_label", []) or []:
        label = lab.get("label", "")
        # Idempotent re-run merges 'copied' (this run) with
        # 'skipped_already_exists' (already in place) — both mean the folder
        # is currently present in the delivery tree.
        present = list(lab.get("copied", []) or []) + list(lab.get("skipped_already_exists", []) or [])
        if present:
            lines.append(f"- 📥 **{label}**: {len(present)} 件取り込み済み")
            for name in present:
                lines.append(f"  - {name}")
        elif lab.get("status") == "no_match":
            lines.append(f"- ℹ {label}: 納品月に該当するフォルダはお客さま環境に存在しませんでした")
    lines.append("")

    # ---- 4. 議事録の取り込み
    s4 = by_step.get(4, {})
    lines.append("### 4. 議事録の取り込み（顧客環境 → 議事録配下）")
    lines.append("各会議フォルダ直下の議事録ファイルを判定し、`議事録/{作業部会, 進捗会議}/` 配下に配置しました。**確定版（ファイル名に `v1.00` を含むもの）のみ**を採用し、未確定の場合は説明テキストを配置します。")
    lines.append("")
    lines.append("また、**前回納品で「※確認中のため次月納品」となっていた会議**については、お客さま環境を再確認し、確定版が出ていれば当月納品に繰り越して取り込んでいます。")
    lines.append("")
    for lab in s4.get("per_label", []) or []:
        label = lab.get("label", "")
        results_l = lab.get("results", []) or []
        copied = [x for x in results_l if x.get("action") == "copied"]
        ph = [x for x in results_l if x.get("action") == "placeholder"]
        carried = [x for x in results_l if x.get("action") == "carried_over"]
        carried_pending = [x for x in results_l if x.get("action") == "carried_over_still_pending"]
        carried_missing = [x for x in results_l if x.get("action") == "carried_over_source_missing"]
        line = f"- **{label}**: 当月分 確定版 {len(copied)} 件 / 当月分 未確定 {len(ph)} 件"
        if carried or carried_pending or carried_missing:
            line += f" / 前回繰越 確定版 {len(carried)} 件"
            if carried_pending:
                line += f" / 前回繰越 まだ未確定 {len(carried_pending)} 件"
            if carried_missing:
                line += f" / 前回繰越 元フォルダなし {len(carried_missing)} 件"
        lines.append(line)
        for x in copied:
            files = ", ".join(x.get("files", []))
            lines.append(f"  - 📥 {x.get('meeting', '')} ← `{files}`")
        for x in carried:
            files = ", ".join(x.get("files", []))
            lines.append(f"  - 🔁 [前回繰越] {x.get('meeting', '')} ← `{files}`（前回はプレースホルダだったが、今回確定版を取り込み）")
        for x in ph:
            reason = x.get("reason", "")
            human_reason = "議事録ファイルがまだ存在しない" if "no minutes file" in reason else f"確定版（v1.00）が存在しない（最新は: {', '.join(x.get('non_final_seen', []))}）"
            lines.append(f"  - ⚠ {x.get('meeting', '')}: {human_reason}")
            lines.append(f"    → `{x.get('placeholder', '')}` を配置しました（来月の納品で改めて確定版を提出してください）")
        for x in carried_pending:
            non_final = ", ".join(x.get("non_final_seen", []) or [])
            lines.append(f"  - ⚠ [前回繰越中] {x.get('meeting', '')}: 今月もまだ確定版が存在しません" + (f"（最新: {non_final}）" if non_final else ""))
            lines.append(f"    → 前回納品フォルダ側のプレースホルダはそのまま残ります。来月再確認します。")
        for x in carried_missing:
            lines.append(f"  - ⚠ [前回繰越中] {x.get('meeting', '')}: お客さま環境から元フォルダが見つかりませんでした")
    lines.append("")

    # ---- 5. 管理簿の取り込み
    s5 = by_step.get(5, {})
    lines.append("### 5. 管理簿の取り込み（3 ファイル）")
    lines.append("お客さま環境（`B2B_導入・構築` / `B2B_連絡チーム`）から管理簿の最新版を読み取り、ファイル名末尾に納品月の月末日を付けて配置しました。")
    lines.append("")
    for x in s5.get("results", []) or []:
        dst_name = Path(x.get("dst", "")).name
        stem = x.get("stem", "")
        if "リスク管理簿" in stem:
            bucket = "プロジェクト管理 / リスク管理表"
        elif "課題・ToDo管理簿" in stem:
            bucket = "プロジェクト管理 / 課題管理表"
        elif "サービス横断会議_連絡台帳" in stem:
            bucket = "本稼働前対応 / 事業者間質問管理表"
        else:
            bucket = ""
        status = x.get("status")
        # Both 'copied' (this run) and 'exists' (idempotent re-run) mean the
        # ledger is now in the delivery tree.
        icon = "📥" if status in ("copied", "exists") else "⚠"
        note = ""
        if status == "missing":
            note = "（顧客環境に元ファイルが見つかりませんでした）"
        lines.append(f"- {icon} **{bucket}** ← `{dst_name}` {note}".rstrip())
    lines.append("")

    # ---- 6. 稼働状況報告書
    s6 = by_step.get(6, {})
    lines.append("### 6. 稼働状況報告書の取り込み（前月度フォルダ）")
    lines.append("稼働状況報告書は **納品月の前月分** を対象とします（例: 4月納品 → 3月度フォルダ）。")
    lines.append("")
    if s6.get("status") == "ok":
        lines.append(f"- 納品月 {s6.get('delivery_month')} → 対象月度 **{s6.get('target_month')}**")
        # Idempotent re-run: 'copied' is what was added this run; existing
        # folders show up under 'skipped_already_exists'. Both mean it's
        # currently in the delivery tree.
        present_s6 = list(s6.get("copied", []) or []) + list(s6.get("skipped_already_exists", []) or [])
        for name in present_s6:
            lines.append(f"- 📥 `{name}/` フォルダを取り込み済み")
    elif s6.get("status") == "no_match":
        lines.append(f"- ⚠ 納品月 {s6.get('delivery_month')} の対象月度 {s6.get('target_month')} に該当するフォルダがお客さま環境に見つかりませんでした")
    elif s6.get("status") == "skipped":
        lines.append(f"- ⚠ ソースフォルダにアクセスできませんでした: {s6.get('reason', '')}")
    lines.append("")

    # ---- 7. 検査仕様書
    s7 = by_step.get(7, {})
    lines.append("### 7. 検査仕様書（xlsx）の生成")
    lines.append(f"📝 `{Path(s7.get('path', '')).name}` を作成済み。")
    sheet_renamed = s7.get("sheet_renamed")
    if sheet_renamed:
        before, after = sheet_renamed.split(" -> ")
        lines.append(f"- メインシート名を「{before}」→「**{after}**」に変更しました")
    else:
        # Idempotent re-run: sheet rename ran on a previous run; report the
        # final name from params.
        lines.append(f"- メインシート名は **第{p.get('n', '')}回中間分** に設定済み")
    lines.append("- セルの色・枠線・フォント・印刷設定など、テンプレートの書式は完全に保持されています（zip+XML 直接編集による安全な書換）")
    lines.append("")

    # ---- 8. 業務完了報告
    s8 = by_step.get(8, {})
    lines.append("### 8. 業務完了報告（docx）の生成")
    lines.append(f"📝 `{Path(s8.get('path', '')).name}` を作成済み。")
    date_replaced8 = s8.get("date_replaced")
    if not date_replaced8:
        date_replaced8 = (
            f"令和{p.get('reiwa_month_end', '')}年"
            f"{p.get('month', '')}月"
            f"{int(str(p.get('month_end_str', '00'))[-2:])}日"
        )
    contract_replaced = s8.get("contract_replaced") or (
        f"令和{p.get('reiwa_fiscal', '')}年度　第{p.get('n', '')}回中間分業務"
    )
    lines.append(f"- 右上の日付: **{date_replaced8.strip()}**")
    lines.append(f"- 履行内容欄: **{contract_replaced}**")
    lines.append("")

    # ---- 9. 統合PF請求書
    s9 = by_step.get(9, {})
    lines.append("### 9. 統合PF請求書（docx）の生成")
    lines.append(f"📝 `{Path(s9.get('path', '')).name}` を作成済み。")
    date_replaced9 = s9.get("date_replaced")
    if not date_replaced9:
        # next_month15_iso is e.g. "2026-05-15" → reconstruct 和暦.
        iso = p.get("next_month15_iso", "")
        if iso and "-" in iso:
            _, m_part, d_part = iso.split("-")
            date_replaced9 = (
                f"　令和{p.get('reiwa_next_month15', '')}年　{int(m_part)}月　{int(d_part)}日"
            )
        else:
            date_replaced9 = ""
    if date_replaced9:
        lines.append(f"- 右上の請求日付: **{date_replaced9.strip()}**（納品月の翌月15日）")
    lines.append("")

    # ---- 次に対応してください
    lines.append("## 次に対応してください")
    lines.append("")
    lines.append("以下は人手での確認・補足が必要な項目です。")
    lines.append("")
    lines.append("**1. Office 文書の目視確認**")
    lines.append("")
    lines.append("以下3ファイルを Excel / Word で開き、レイアウト・書式・置換結果に問題がないか確認してください。")
    lines.append("")
    lines.append("- 検査仕様書（シート名が `第" + str(p.get("n", "")) + "回中間分` になっているか）")
    lines.append("- 業務完了報告（日付・履行内容）")
    lines.append("- 統合PF請求書（右上の日付）")
    lines.append("")
    if placeholder_entries:
        lines.append("**2. 議事録の未確定項目**")
        lines.append("")
        lines.append("以下の議事録は、お客さま環境にまだ確定版（v1.00）が存在しないため、暫定的に `※確認中のため次月納品.txt` を配置しています。来月の納品時に確定版を改めて差し替えてください。")
        lines.append("")
        for label, meeting in placeholder_entries:
            lines.append(f"- {label}: **{meeting}**")
        lines.append("")
    lines.append("**3. 各管理簿の最新性確認**")
    lines.append("")
    lines.append("管理簿（リスク管理簿 / 課題・ToDo管理簿 / サービス横断会議連絡台帳）は実行時点のお客さま環境最新版を取り込んでいます。納品直前にお客さま側で更新が入っていないか念のためご確認ください。")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("**生成元スキル**: `~/.claude/skills/delivery-folder-creator/`")
    lines.append("")
    return "\n".join(lines)


def write_summary_md(params: dict, results: list[dict]) -> Path:
    """Write the summary md to <root>/納品作成サマリ_<YYYYMMDD>.md."""
    out = Path(params["root_dir"]) / f"納品作成サマリ_{params['today_str']}.md"
    out.write_text(format_summary_md(params, results), encoding="utf-8")
    return out


# ---- CLI ------------------------------------------------------------------
def main() -> int:
    p = argparse.ArgumentParser(description="Create monthly delivery folder for 大阪市統合PF")
    p.add_argument(
        "action",
        choices=["compute", "step1", "step2", "step3", "step4", "step5",
                 "step6", "step7", "step8", "step9", "all", "summary"],
    )
    p.add_argument("--month", required=True, help="Delivery month (e.g. '2026-05', '5月', '202605')")
    p.add_argument("--base", required=True, help="Base delivery folder")
    p.add_argument("--today", default=None, help="Override today (YYYYMMDD), for testing")
    p.add_argument("--n", type=int, default=None, help="Override n, for testing")
    args = p.parse_args()

    base = Path(args.base)
    params = compute_params(args.month, base, today_str=args.today, n_override=args.n)

    actions = {
        "compute": lambda: {"step": 0, "status": "ok", **params},
        "step1": lambda: step1_create_root(params),
        "step2": lambda: step2_create_subfolders(params),
        "step3": lambda: step3_copy_meetings(params),
        "step4": lambda: step4_copy_minutes(params),
        "step5": lambda: step5_copy_management_files(params),
        "step6": lambda: step6_copy_kasou_joukyou(params),
        "step7": lambda: step7_create_kensa(params),
        "step8": lambda: step8_create_gyomu(params),
        "step9": lambda: step9_create_seikyuu(params),
    }

    if args.action in ("all", "summary"):
        # `all` and `summary` share the same execution: step1〜9 are idempotent,
        # so re-running `summary` after individual stepN calls is safe and just
        # ensures the summary md exists (no destructive side-effects).
        results = [actions["compute"]()]
        for k in ("step1", "step2", "step3", "step4", "step5",
                  "step6", "step7", "step8", "step9"):
            results.append(actions[k]())
        # Write a single Markdown summary alongside the deliverables.
        summary_path = write_summary_md(params, results)
        results.append({"step": 10, "status": "ok",
                        "summary_md": str(summary_path),
                        "mode": args.action})
        print(json.dumps(results, ensure_ascii=False, indent=2))
        return 0

    result = actions[args.action]()
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
