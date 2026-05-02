"""Inspect the three Office templates and print where date / sheet-name / contract text appear.

Run once during skill development to determine the exact runs/cells we must update,
so create_delivery_folder.py knows what to replace without touching formatting.
"""
from __future__ import annotations
import sys
from pathlib import Path

import openpyxl
import docx

BASE = Path(__file__).resolve().parents[1] / "assets" / "templates"
KENSA = BASE / "kensa_yoshiki.xlsx"
GYOMU = BASE / "gyomu_kanryou.docx"
SEIKYUU = BASE / "seikyuu.docx"


def inspect_xlsx(path: Path) -> None:
    print(f"\n=== XLSX: {path.name} ===")
    wb = openpyxl.load_workbook(path)
    print("Sheets:", wb.sheetnames)
    for sn in wb.sheetnames:
        ws = wb[sn]
        print(f"  -- '{sn}' ({ws.max_row}x{ws.max_column}) --")
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 25)):
            for cell in row:
                if cell.value is not None:
                    v = str(cell.value).replace("\n", "\\n")
                    if len(v) > 80:
                        v = v[:80] + "..."
                    print(f"    {cell.coordinate}: {v!r}")


def inspect_docx(path: Path) -> None:
    print(f"\n=== DOCX: {path.name} ===")
    doc = docx.Document(path)

    print("--- Paragraphs (top-level) ---")
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        runs = [(r.text, r.font.name, r.font.size.pt if r.font.size else None) for r in p.runs]
        print(f"[P{i}] align={p.alignment} text={p.text!r}")
        for j, (t, fn, fs) in enumerate(runs):
            print(f"     run{j}: font={fn} size={fs} text={t!r}")

    print("--- Tables ---")
    for ti, t in enumerate(doc.tables):
        print(f"  -- Table {ti}: {len(t.rows)}x{len(t.columns)} --")
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip().replace("\n", "\\n")
                if not text:
                    continue
                print(f"    R{ri}C{ci}: {text!r}")
                for pi, p in enumerate(cell.paragraphs):
                    if not p.text.strip():
                        continue
                    for j, r in enumerate(p.runs):
                        if r.text:
                            print(f"      P{pi}.run{j}: {r.text!r} (font={r.font.name}, size={r.font.size.pt if r.font.size else None})")


def main() -> int:
    inspect_xlsx(KENSA)
    inspect_docx(GYOMU)
    inspect_docx(SEIKYUU)
    return 0


if __name__ == "__main__":
    sys.exit(main())
