#!/usr/bin/env python3
"""
generate_minutes.py — テンプレートベースの議事録 .docx 生成スクリプト

入力: meeting_data.json + assets/template.docx
出力: 議事録 .docx
"""
import argparse
import json
import os
import sys
import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Emu, Twips, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


SKILL_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = SKILL_DIR / "assets" / "template.docx"
FONT_NAME = "Meiryo UI"
FONT_SIZE = Pt(10)

# ── indent constants (EMU) from model answer ──
SECTION_HEADING_LEFT = -1270
SECTION_HEADING_FIRST = 1270
LEVEL0_LEFT = 0                  # agenda title
LEVEL1_LEFT = 360045             # ● new statement (hanging)
LEVEL1_FIRST = -131445
LEVEL2_LEFT = 540385             # → response to level 1 (hanging)
LEVEL2_FIRST = -180340
LEVEL3_LEFT = LEVEL2_LEFT        # → same indent as level 2 (same symbol)
LEVEL3_FIRST = LEVEL2_FIRST
SUBTOPIC_LEFT = 419100           # sub-topic title
NEXT_MTG_LEFT = 127000           # next meeting info

# ── table column widths (twips/dxa) from model answer ──
DECISION_TABLE_WIDTH = 9480
DECISION_COL_NO = 600
DECISION_COL_TEXT = 8880

TODO_TABLE_WIDTH = 9480
TODO_COL_NO = 600
TODO_COL_TODO = 5640
TODO_COL_OWNER = 1680
TODO_COL_DUE = 1560

# ── tag color (RGB hex) ──
TAG_COLOR = "FF0000"  # red

# ── bullet numId (will be created at runtime) ──
BULLET_NUM_ID = "100"


def set_font(run, bold=False, size=None, color=None):
    """Apply Meiryo UI font to a run."""
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = size or FONT_SIZE
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_indent(paragraph, left=None, first_line=None):
    """Set paragraph indent in EMU."""
    pf = paragraph.paragraph_format
    if left is not None:
        pf.left_indent = left
    if first_line is not None:
        pf.first_line_indent = first_line


def _add_numPr(paragraph, num_id, ilvl):
    """Add numbering properties (bullet) to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), str(ilvl))
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_elem)
    numPr.append(numId_elem)
    pPr.append(numPr)


def _ensure_bullet_numbering(doc):
    """Create bullet numbering definitions in the document's numbering part.

    Creates abstractNum/num with 3 ilvl:
      ilvl=0: ● (Symbol \\uf0b7) — for level 1 (new statement/topic)
      ilvl=1: → (Wingdings \\uf0e0) — for level 2 (response)
      ilvl=2: → (Wingdings \\uf0e0) — for level 3 (deeper response)
    """
    try:
        numbering_elem = doc.part.numbering_part._element
    except (KeyError, AttributeError):
        # Template should have a numbering part; if not, bail out gracefully
        return

    # Check if already created
    for n in numbering_elem.findall(qn("w:num")):
        if n.get(qn("w:numId")) == BULLET_NUM_ID:
            return

    # abstractNum
    absNum = OxmlElement("w:abstractNum")
    absNum.set(qn("w:abstractNumId"), BULLET_NUM_ID)
    mlt = OxmlElement("w:multiLevelType")
    mlt.set(qn("w:val"), "hybridMultilevel")
    absNum.append(mlt)

    # ilvl 0: bullet ● (Symbol font)
    absNum.append(_make_bullet_lvl("0", "\uf0b7", "Symbol", "720", "360"))
    # ilvl 1: arrow → (Wingdings font)
    absNum.append(_make_bullet_lvl("1", "\uf0e0", "Wingdings", "1140", "360"))
    # ilvl 2: arrow → deeper (Wingdings font)
    absNum.append(_make_bullet_lvl("2", "\uf0e0", "Wingdings", "1560", "360"))

    # OOXML requires w:abstractNum before all w:num elements
    first_num = numbering_elem.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(absNum)
    else:
        numbering_elem.append(absNum)

    # num → abstractNum reference
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), BULLET_NUM_ID)
    absNumId = OxmlElement("w:abstractNumId")
    absNumId.set(qn("w:val"), BULLET_NUM_ID)
    num.append(absNumId)
    numbering_elem.append(num)


def _make_bullet_lvl(ilvl, char, font, left_twips, hanging_twips):
    """Create a <w:lvl> element for a bullet level."""
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), ilvl)

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "bullet")
    lvl.append(numFmt)

    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), char)
    lvl.append(lvlText)

    lvlJc = OxmlElement("w:lvlJc")
    lvlJc.set(qn("w:val"), "left")
    lvl.append(lvlJc)

    # paragraph properties (indent)
    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), left_twips)
    ind.set(qn("w:hanging"), hanging_twips)
    pPr.append(ind)
    lvl.append(pPr)

    # run properties (font for bullet char)
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:hint"), "default")
    rPr.append(rFonts)
    lvl.append(rPr)

    return lvl


def add_paragraph(doc, text="", bold=False, left=None, first_line=None,
                  font_size=None, num_id=None, ilvl=None):
    """Add a paragraph with Meiryo UI font and optional indent/bullet."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, size=font_size)
    set_paragraph_indent(p, left=left, first_line=first_line)
    if num_id is not None and ilvl is not None:
        _add_numPr(p, num_id, ilvl)
    return p


def add_section_heading(doc, title):
    """Add a section heading like <アジェンダ>, <決定事項>, etc."""
    p = add_paragraph(doc, title, bold=True,
                      left=SECTION_HEADING_LEFT, first_line=SECTION_HEADING_FIRST)
    return p


def add_mixed_paragraph(doc, parts, left=None, first_line=None,
                        num_id=None, ilvl=None):
    """Add a paragraph with mixed formatting runs.
    parts: list of (text, bold, color_or_None) tuples.
    Also accepts legacy (text, bold) tuples.
    """
    p = doc.add_paragraph()
    set_paragraph_indent(p, left=left, first_line=first_line)
    for item in parts:
        if len(item) == 3:
            text, bold, color = item
        else:
            text, bold = item
            color = None
        run = p.add_run(text)
        set_font(run, bold=bold, color=color)
    if num_id is not None and ilvl is not None:
        _add_numPr(p, num_id, ilvl)
    return p


def build_header_table(doc, data):
    """Fill the header table (author/approver/date/title) in the Word header area.

    Template header has a 3-row x 2-col table:
      [0,0] 作成者: (Name)    [0,1] 作成日: YYYY/MM/DD
      [1,0] 承認者: (Name)    [1,1] 承認日: YYYY/MM/DD
      [2,0] ミーティング議事録  (colspan=2, bold+underline, center)
    """
    header_info = data.get("header")
    if not header_info:
        return

    section = doc.sections[0]
    header = section.header
    # Find the table in the header
    header_tbls = header._element.findall(qn("w:tbl"))
    if not header_tbls:
        return

    tbl_elem = header_tbls[0]
    rows = tbl_elem.findall(qn("w:tr"))
    if len(rows) < 3:
        return

    def _replace_header_cell(row_elem, col_idx, new_text):
        """Replace text in a header table cell while preserving formatting."""
        cells = row_elem.findall(qn("w:tc"))
        if col_idx >= len(cells):
            return
        cell = cells[col_idx]
        # Find all w:t elements and replace their text
        t_elems = cell.findall(".//" + qn("w:t"))
        if t_elems:
            # Clear all but first, set first to new text
            t_elems[0].text = new_text
            for t in t_elems[1:]:
                t.text = ""

    author = header_info.get("author", "")
    approver = header_info.get("approver", "")
    created_date = header_info.get("created_date", "")
    approved_date = header_info.get("approved_date", "")
    doc_title = header_info.get("doc_title", "ミーティング議事録")

    _replace_header_cell(rows[0], 0, f"作成者 : {author}")
    _replace_header_cell(rows[0], 1, f"作成日: {created_date}")
    _replace_header_cell(rows[1], 0, f"承認者 : {approver}")
    _replace_header_cell(rows[1], 1, f"承認日: {approved_date}")
    _replace_header_cell(rows[2], 0, doc_title)


def build_info_table(doc, data):
    """Fill the pre-existing basic info table (5x2) from template."""
    table = doc.tables[0]
    rows = table.rows
    # Row 0: datetime
    _set_cell_text(rows[0].cells[1], data.get("datetime", ""))
    # Row 1: place
    _set_cell_text(rows[1].cells[1], data.get("place", ""))
    # Row 2: title
    _set_cell_text(rows[2].cells[1], data.get("meeting_title", ""))
    # Row 3: participants
    _set_cell_text(rows[3].cells[1], _format_participants(data.get("participants", [])))
    # Row 4: materials
    _set_cell_text(rows[4].cells[1], data.get("materials", ""))


def _add_text_with_breaks(paragraph, text, bold=False, color=None):
    """Add text to a paragraph, converting \\n to Word line breaks (<w:br/>)."""
    parts = text.split("\n")
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        set_font(run, bold=bold, color=color)
        if i < len(parts) - 1:
            run._element.append(OxmlElement("w:br"))


def _set_cell_text(cell, text):
    """Set cell text with Meiryo UI font. Handles \\n as Word line breaks."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.clear()
    cell.paragraphs[0].clear()
    # Remove extra paragraphs (from previous content)
    while len(cell.paragraphs) > 1:
        cell._tc.remove(cell.paragraphs[-1]._p)
    _add_text_with_breaks(cell.paragraphs[0], text)


def _format_participants(participants):
    """Format participant list.
    participants: list of {"org": "UM", "members": "中山、小山（記）"}
    """
    lines = []
    for p in participants:
        org = p.get("org", "")
        members = p.get("members", "")
        lines.append(f"{org}）{members}")
    return "\n".join(lines)


def build_agenda_section(doc, agenda_items):
    """Populate the <アジェンダ> items with numbered paragraphs (1., 2., ...)."""
    for i, item in enumerate(agenda_items, 1):
        add_paragraph(doc, f"{i}. {item}", left=0)


def build_discussion_section(doc, discussions):
    """Build <議事詳細> section with proper indent hierarchy and bullets.

    Bullet levels (numId=100):
      ilvl=0: ● — for level 1 (new statement/topic)
      ilvl=1: → — for level 2 (response)
      ilvl=2: → — for level 3 (deeper response)
    """
    nid = int(BULLET_NUM_ID)

    for topic_idx, topic in enumerate(discussions, 1):
        # Agenda title (level 0) - numbered bold title + optional non-bold note
        title = topic.get("title", "")
        numbered_title = f"{topic_idx}. {title}"
        note = topic.get("title_note", "")
        if note:
            parts = [(numbered_title, True, None), ("\u3000" + note, False, None)]
            add_mixed_paragraph(doc, parts, left=LEVEL0_LEFT)
        else:
            add_paragraph(doc, numbered_title, bold=True, left=LEVEL0_LEFT)

        for stmt in topic.get("statements", []):
            level = stmt.get("level", 2)
            text = stmt.get("text", "")
            speaker = stmt.get("speaker", "")
            tag = stmt.get("tag", "")

            # Build speaker suffix
            speaker_suffix = f"（{speaker}）" if speaker else ""

            if level == 1:
                # Overview statement — bullet ● (ilvl=0), same indent as level 2
                full_text = text + speaker_suffix
                if tag:
                    parts = [(full_text, False, None), (tag, True, TAG_COLOR)]
                    add_mixed_paragraph(doc, parts,
                                        left=LEVEL1_LEFT, first_line=LEVEL1_FIRST,
                                        num_id=nid, ilvl=0)
                else:
                    add_paragraph(doc, full_text,
                                  left=LEVEL1_LEFT, first_line=LEVEL1_FIRST,
                                  num_id=nid, ilvl=0)
            elif level == 2:
                # Response — bullet → (ilvl=1) + hanging indent
                if tag:
                    parts = [(text + speaker_suffix, False, None),
                             (tag, True, TAG_COLOR)]
                    add_mixed_paragraph(doc, parts,
                                        left=LEVEL2_LEFT, first_line=LEVEL2_FIRST,
                                        num_id=nid, ilvl=1)
                else:
                    add_paragraph(doc, text + speaker_suffix,
                                  left=LEVEL2_LEFT, first_line=LEVEL2_FIRST,
                                  num_id=nid, ilvl=1)
            elif level == 3:
                # Deeper response — bullet → (ilvl=2) + deeper hanging indent
                if tag:
                    parts = [(text + speaker_suffix, False, None),
                             (tag, True, TAG_COLOR)]
                    add_mixed_paragraph(doc, parts,
                                        left=LEVEL3_LEFT, first_line=LEVEL3_FIRST,
                                        num_id=nid, ilvl=2)
                else:
                    add_paragraph(doc, text + speaker_suffix,
                                  left=LEVEL3_LEFT, first_line=LEVEL3_FIRST,
                                  num_id=nid, ilvl=2)
            elif level == "subtopic":
                # Sub-topic title
                add_paragraph(doc, text + (f" {note}" if (note := stmt.get("note", "")) else ""),
                              left=SUBTOPIC_LEFT)

        # Blank line between topics
        add_paragraph(doc, "")


def build_next_meeting(doc, next_meeting):
    """Build <次回ミーティング> section."""
    if not next_meeting:
        return
    dt = next_meeting.get("datetime", "")
    place = next_meeting.get("place", "")
    attendees = next_meeting.get("attendees", [])

    add_paragraph(doc, f"日時:\t{dt}", left=NEXT_MTG_LEFT)
    add_paragraph(doc, f"場所:\t{place}", left=NEXT_MTG_LEFT)

    if attendees:
        first = attendees[0]
        add_paragraph(doc, f"出席者:\t{first.get('org', '')}）{first.get('members', '')}",
                      left=NEXT_MTG_LEFT)
        for att in attendees[1:]:
            add_paragraph(doc, f"\t\t{att.get('org', '')}）{att.get('members', '')}",
                          left=NEXT_MTG_LEFT)


def clear_template_body(doc):
    """Remove all paragraphs/tables after the first table (keep header table + sectPr)."""
    body = doc.element.body
    # Find the first table
    tables = body.findall(qn("w:tbl"))
    if not tables:
        return

    # Keep the first table (basic info), remove everything after it
    # BUT preserve w:sectPr (section properties) — removing it breaks doc.add_table()
    first_table = tables[0]
    found_first = False
    to_remove = []
    for child in body:
        if child is first_table:
            found_first = True
            continue
        if found_first and child.tag != qn("w:sectPr"):
            to_remove.append(child)
    for elem in to_remove:
        body.remove(elem)


def _set_table_grid(tbl, col_widths):
    """Set tblGrid column widths (twips/dxa) and total table width."""
    # Remove existing tblGrid
    for old_grid in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old_grid)

    grid = OxmlElement("w:tblGrid")
    for w in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    # Insert after tblPr
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(grid)
    else:
        tbl.insert(0, grid)

    # Set total table width
    total = sum(col_widths)
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    # Remove existing tblW
    for old_w in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old_w)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _set_cell_width(cell, width_twips):
    """Set explicit cell width (tcW) in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing tcW
    for old_w in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old_w)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _validate_data(data):
    """Basic validation of input JSON structure."""
    if not isinstance(data, dict):
        raise ValueError("入力データがJSON objectではありません")
    required_keys = ["meeting_title", "datetime", "participants", "discussions"]
    missing = [k for k in required_keys if k not in data]
    if missing:
        raise ValueError(f"必須キーが不足しています: {', '.join(missing)}")
    if not isinstance(data.get("participants", []), list):
        raise ValueError("participants はリストである必要があります")
    if not isinstance(data.get("discussions", []), list):
        raise ValueError("discussions はリストである必要があります")
    for i, todo in enumerate(data.get("todos", [])):
        if not isinstance(todo, dict):
            raise ValueError(f"todos[{i}] がオブジェクトではありません")


def generate(data, output_path):
    """Main generation function."""
    _validate_data(data)
    doc = Document(str(TEMPLATE_PATH))

    # Ensure bullet numbering definitions exist
    _ensure_bullet_numbering(doc)

    # Fill header table (author/approver/date/title)
    build_header_table(doc, data)

    # Fill basic info table
    build_info_table(doc, data)

    # Clear template body below first table
    clear_template_body(doc)

    # <アジェンダ>
    add_section_heading(doc, "<アジェンダ>")
    build_agenda_section(doc, data.get("agenda", []))
    add_paragraph(doc, "")

    # <決定事項>
    add_section_heading(doc, "<決定事項>")
    # Need to create decisions table
    _create_decisions_table(doc, data.get("decisions", []))
    add_paragraph(doc, "")

    # <ToDos>
    add_section_heading(doc, "<ToDos>")
    _create_todos_table(doc, data.get("todos", []))
    add_paragraph(doc, "")

    # <議事詳細>
    add_section_heading(doc, "<議事詳細>")
    build_discussion_section(doc, data.get("discussions", []))

    # <次回ミーティング>
    add_section_heading(doc, "<次回ミーティング>")
    build_next_meeting(doc, data.get("next_meeting"))
    add_paragraph(doc, "")

    # 以上
    add_paragraph(doc, "以上")

    doc.save(str(output_path))
    print(f"Generated: {output_path}")


def _create_decisions_table(doc, decisions):
    """Create a new decisions table matching model answer style.

    Column widths (dxa/twips): No=600, 決定事項=8880, total=9480
    """
    col_widths = [DECISION_COL_NO, DECISION_COL_TEXT]
    table = doc.add_table(rows=1, cols=2)
    tbl = table._tbl

    # Ensure tblPr exists
    if tbl.tblPr is None:
        tbl.insert(0, OxmlElement("w:tblPr"))

    _set_table_grid(tbl, col_widths)
    _set_table_borders(tbl)

    # Header row
    hdr = table.rows[0]
    _set_cell_text(hdr.cells[0], "No")
    _set_cell_text(hdr.cells[1], "決定事項")
    _set_cell_width(hdr.cells[0], col_widths[0])
    _set_cell_width(hdr.cells[1], col_widths[1])
    _shade_row(hdr)

    # Data rows
    for i, decision in enumerate(decisions, 1):
        row = table.add_row()
        _set_cell_text(row.cells[0], str(i))
        _set_cell_text(row.cells[1], decision)
        _set_cell_width(row.cells[0], col_widths[0])
        _set_cell_width(row.cells[1], col_widths[1])


def _create_todos_table(doc, todos):
    """Create a new ToDo table matching model answer style.

    Column widths (dxa/twips): No=600, ToDo=5640, 担当者=1680, 期限=1560, total=9480
    """
    col_widths = [TODO_COL_NO, TODO_COL_TODO, TODO_COL_OWNER, TODO_COL_DUE]
    table = doc.add_table(rows=1, cols=4)
    tbl = table._tbl

    # Ensure tblPr exists
    if tbl.tblPr is None:
        tbl.insert(0, OxmlElement("w:tblPr"))

    _set_table_grid(tbl, col_widths)
    _set_table_borders(tbl)

    # Header row
    hdr = table.rows[0]
    headers = ["No", "To Do", "担当者", "期限"]
    for ci, (cell, header) in enumerate(zip(hdr.cells, headers)):
        _set_cell_text(cell, header)
        _set_cell_width(cell, col_widths[ci])
    _shade_row(hdr)

    # Data rows
    for i, todo in enumerate(todos, 1):
        row = table.add_row()
        values = [f"{i}.", todo.get("action", ""),
                  todo.get("owner", ""), todo.get("due", "")]
        for ci, (cell, val) in enumerate(zip(row.cells, values)):
            _set_cell_text(cell, val)
            _set_cell_width(cell, col_widths[ci])


def _set_table_borders(tbl):
    """Set table borders to #999999 6pt."""
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "999999")
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)


def _shade_row(row):
    """Apply pct20 shading to a row (header)."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "pct20")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "auto")
        tcPr.append(shd)


def main():
    parser = argparse.ArgumentParser(description="Generate meeting minutes .docx")
    parser.add_argument("--data", required=True, help="Path to meeting_data.json")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate(data, args.output)


if __name__ == "__main__":
    main()
