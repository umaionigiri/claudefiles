#!/usr/bin/env python3
"""
fix_minutes.py — 議事録の表現ルール準拠自動修正

入力: 対象の議事録 .docx
出力: 修正済み .docx + 修正箇所のサマリー（stdout）

自動修正範囲:
- 口語表現→ビジネス表現の置換
- 敬体→常体の変換
- 「させて頂く」「頂く」の除去
- フォントの Meiryo UI 統一
- 「～を行う」→簡潔な動詞への置換
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

FONT_NAME = "Meiryo UI"
FONT_SIZE = Pt(10)

# ── 表現置換ルール ──
# (pattern, replacement) — 順序が重要（長い方を先にマッチ）
EXPRESSION_RULES = [
    # させて頂く系
    (r"させて頂きます", "します"),  # intermediate step
    (r"させていただきます", "します"),
    (r"させて頂く", "する"),
    (r"させていただく", "する"),
    (r"させて頂いた", "した"),
    (r"させていただいた", "した"),

    # 頂く系（動詞の後の「頂く」→除去）
    (r"確認頂いた", "確認した"),
    (r"ご確認頂いた", "確認した"),
    (r"ご確認いただいた", "確認した"),
    (r"ご回答頂いた", "回答した"),
    (r"ご回答いただいた", "回答した"),
    (r"ご連絡頂いた", "連絡した"),
    (r"ご連絡いただいた", "連絡した"),
    (r"送付頂いた", "送付した"),
    (r"送付いただいた", "送付した"),

    # 「を行う」系
    (r"検討を行う", "検討する"),
    (r"検討を行った", "検討した"),
    (r"確認を行う", "確認する"),
    (r"確認を行った", "確認した"),
    (r"調査を行う", "調査する"),
    (r"調査を行った", "調査した"),
    (r"実施を行う", "実施する"),
    (r"対応を行う", "対応する"),
    (r"対応を行った", "対応した"),
    (r"作成を行う", "作成する"),
    (r"作成を行った", "作成した"),
    (r"報告を行う", "報告する"),
    (r"報告を行った", "報告した"),
    (r"説明を行う", "説明する"),
    (r"説明を行った", "説明した"),

    # 口語→書き言葉
    (r"やる(?=[。、\s])", "実施する"),
    (r"やった(?=[。、\s])", "実施した"),
    (r"サマる", "要約する"),
    (r"サマった", "要約した"),
    (r"ケアする", "確認する"),
    (r"リダンダント", "冗長"),

    # ～していく系
    (r"していきたい(?=[。、\s])", "したい"),
    (r"していく(?=[。、\s])", "する"),
    (r"してまいります", "する"),
    (r"してまいりたい", "したい"),

    # 「夏休み」→「夏季休暇」
    (r"夏休み", "夏季休暇"),

    # 「バケーション」→「休暇」
    (r"バケーション", "休暇"),

    # 「～ので」→「～ため」（自動修正可能なパターン）
    (r"ないので([、,])", r"ないため\1"),
    (r"あるので([、,])", r"あるため\1"),
    (r"いるので([、,])", r"いるため\1"),
    (r"したので([、,])", r"したため\1"),
    (r"れるので([、,])", r"れるため\1"),
    (r"えるので([、,])", r"えるため\1"),
    (r"するので([、,])", r"するため\1"),
    (r"ないなので([、,])", r"ないため\1"),

    # 矛盾した擬似熟語
    (r"基本概念", "概念"),
    (r"詳細概要", "詳細"),
    (r"概要詳細", "概要"),
]

# ── 敬体→常体 変換ルール ──
POLITE_TO_PLAIN = [
    # 順序重要: 長い方を先に
    (r"でございます(?=[。、\s])", "である"),
    (r"いたします(?=[。、\s])", "する"),
    (r"おります(?=[。、\s])", "いる"),
    (r"ございます(?=[。、\s])", "ある"),
    (r"でした(?=[。、\s])", "であった"),
    (r"ました(?=[。、\s])", "た"),
    (r"ません(?=[。、\s])", "ない"),
    (r"です(?=[。、\s])", "である"),
    (r"ます(?=[。、\s])", "る"),
]


class FixLog:
    """Track all fixes applied."""
    def __init__(self):
        self.entries = []

    def add(self, category, original, fixed, location=""):
        self.entries.append({
            "category": category,
            "original": original,
            "fixed": fixed,
            "location": location,
        })

    def summary(self):
        if not self.entries:
            return "修正箇所なし"
        lines = [f"修正箇所: {len(self.entries)}件", ""]
        for i, e in enumerate(self.entries, 1):
            lines.append(f"{i}. [{e['category']}] 「{e['original']}」→「{e['fixed']}」")
            if e["location"]:
                lines.append(f"   場所: {e['location']}")
        return "\n".join(lines)


def fix_text(text, log, location=""):
    """Apply all text-level fixes and return corrected text."""
    original = text

    # Expression rules
    for pattern, replacement in EXPRESSION_RULES:
        match = re.search(pattern, text)
        if match:
            text = re.sub(pattern, replacement, text)
            log.add("表現置換", match.group(), replacement, location)

    # Polite → plain
    for pattern, replacement in POLITE_TO_PLAIN:
        match = re.search(pattern, text)
        if match:
            text = re.sub(pattern, replacement, text)
            log.add("敬体→常体", match.group(), replacement, location)

    return text


def fix_font(run, log, location=""):
    """Ensure Meiryo UI font."""
    changed = False
    if run.font.name and run.font.name != FONT_NAME:
        log.add("フォント", f"font={run.font.name}", f"font={FONT_NAME}", location)
        changed = True
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    return changed


def fix_document(input_path, output_path):
    """Apply all auto-fixable corrections to the document."""
    doc = Document(str(input_path))
    log = FixLog()

    # Fix paragraphs
    for i, p in enumerate(doc.paragraphs):
        location = f"段落{i}"
        for run in p.runs:
            if run.text.strip():
                # Fix font
                fix_font(run, log, location)
                # Fix text
                new_text = fix_text(run.text, log, location)
                if new_text != run.text:
                    run.text = new_text

    # Fix table cells
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                location = f"テーブル{ti} [{ri},{ci}]"
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip():
                            fix_font(run, log, location)
                            new_text = fix_text(run.text, log, location)
                            if new_text != run.text:
                                run.text = new_text

    doc.save(str(output_path))
    print(f"Fixed: {output_path}")
    print(log.summary())
    return log


def main():
    parser = argparse.ArgumentParser(description="Auto-fix meeting minutes .docx")
    parser.add_argument("--input", required=True, help="Input .docx to fix")
    parser.add_argument("--output", required=True, help="Output fixed .docx")
    args = parser.parse_args()

    fix_document(args.input, args.output)


if __name__ == "__main__":
    main()
