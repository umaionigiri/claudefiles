#!/usr/bin/env python3
"""
review_minutes.py — 議事録チェックリストに基づくレビュー

入力: 対象の議事録 .docx
出力: レビュー結果 .docx（指摘事項・改善点一覧）
"""
import argparse
import json
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

FONT_NAME = "Meiryo UI"

# ── 口語表現→ビジネス表現 変換辞書 ──
COLLOQUIAL_MAP = {
    "見ている": "把握している",
    "回す": "運用する",
    "しっかりと": "精緻に",
    "はず": "見込みである",
    "やる": "実施する",
    "夏休み": "夏季休暇",
    "バケーション": "休暇",
    "サマる": "要約する",
    "ケアする": "確認する",
    "リダンダント": "冗長",
    "イメージ": None,  # 文脈依存 — 指摘のみ
    "させて頂く": "する",
    "させていただく": "する",
    "していく": "する",
    "していきたい": "したい",
    "確認していく": "確認する",
    "を行う": "する",
    "を行った": "した",
    "を行い": "し",
    "要検討": None,  # 多用チェック
    "要対応": None,
    "要確認": None,
}

# ── 追加表現チェック（パターン → 指摘メッセージ, 改善案） ──
ADDITIONAL_EXPRESSION_CHECKS = [
    # 「～ので」→「～ため」
    (r"(?:ない|ある|いる|した|ない|れる|える|する)[のな]で[、,]",
     "「～ので」が使われている", "「～ため」に置き換える"),
    # 「本PJ」「当PJ」「本件」「当社」等の多用
    (r"(?:本|当)(?:PJ|プロジェクト|件|案件|システム|業務)",
     "「本～」「当～」が使われている", "固有名詞で書くか「今回のPJでは」等に言い換える"),
    # 矛盾した擬似熟語
    (r"基本概念|詳細概要|概要詳細|簡易詳細",
     "矛盾した擬似熟語が使われている", "「概念」「詳細」「概要」等、正確な語に置き換える"),
    # 「および」の前に読点なし
    (r"[^\s、,]および",
     "「および」の前に読点がない", "「A、およびB」のように読点を付ける。並列は「・」で表記する"),
    # 「鑑みる」の乱用
    (r"(?:を|に)鑑み",
     "「鑑みる」が使われている", "「先例に照らす」意味でなければ「考慮する」「踏まえる」に置き換える"),
    # 「基本的に」単体使用
    (r"基本的に(?!.*(?:例外|ただし|一方|応用))",
     "「基本的に」が単体で使われている", "基本-例外のセットで使うか、「原則として」に言い換える"),
    # 「お」「ご」の不適切な使用（漢語に「お」、和語に「ご」）
    (r"お(?:検討|確認|対応|実施|報告|作成|調査|検証|承認|決裁)",
     "漢語に「お」が付いている", "漢語には「ご」を使う（例: 「お検討」→「ご検討」）。ただし常体の議事録では敬語自体を避ける"),
]

# 時期の曖昧表現パターン
VAGUE_TIME_PATTERNS = [
    r"少し前", r"先だって", r"今後は(?!.*(?:\d{4}|年度))", r"近日中",
    r"そのうち", r"近いうち", r"いずれ",
    r"(?<!\d[/\-])昨日(?!\s*[\(（\d])", r"(?<!\d[/\-])明日(?!\s*[\(（\d])",
    r"(?<!\d[/\-])今週(?!\s*[\(（\d])", r"(?<!\d[/\-])来週(?!\s*[\(（\d])",
]

# 指示代名詞パターン
DEMONSTRATIVE_PATTERNS = [
    r"(?:^|[、。\s])これ(?:[はがをにで、。])",
    r"(?:^|[、。\s])それ(?:[はがをにで、。])",
    r"(?:^|[、。\s])あれ(?:[はがをにで、。])",
    r"この(?!度|たび|ため)",
    r"その(?!ため|結果|後|他|際|場合|上|中|間)",
    r"あの",
    r"そういった",
    r"そういう",
]

# 敬体パターン（議事録は常体であるべき）
POLITE_PATTERNS = [
    (r"です[。\s]", "である。"),
    (r"ます[。\s]", "る。（常体に変換）"),
    (r"でした[。\s]", "であった。"),
    (r"ました[。\s]", "た。（常体に変換）"),
    (r"ません[。\s]", "ない。"),
    (r"でしょう", "であろう"),
]

# 失礼表現パターン
RUDE_PATTERNS = [
    r"レベルが低[いく]",
    r"能力が(?:ない|不足)",
    r"無能",
    r"使えない",
]


class Finding:
    """A single review finding."""
    def __init__(self, category, check_no, text, location, suggestion=""):
        self.category = category
        self.check_no = check_no
        self.text = text
        self.location = location
        self.suggestion = suggestion


def extract_text_paragraphs(doc):
    """Extract all paragraph texts with their index."""
    result = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            result.append((i, p.text.strip()))
    return result


def check_font(doc):
    """Check #1: Font should be Meiryo UI."""
    findings = []
    non_meiryo = 0
    total = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text.strip():
                total += 1
                fname = run.font.name
                if fname and fname != FONT_NAME:
                    non_meiryo += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip():
                            total += 1
                            fname = run.font.name
                            if fname and fname != FONT_NAME:
                                non_meiryo += 1
    if non_meiryo > 0 and total > 0:
        pct = non_meiryo / total * 100
        findings.append(Finding(
            "基本情報", 1,
            f"Meiryo UI 以外のフォントが {non_meiryo}/{total} 箇所 ({pct:.0f}%) で使用されている",
            "文書全体",
            "すべてのテキストを Meiryo UI に統一する"
        ))
    return findings


def check_recorder_mark(doc):
    """Check #3: 作成者に（記）があるか."""
    findings = []
    if len(doc.tables) > 0:
        table = doc.tables[0]
        if len(table.rows) > 3:
            cell_text = table.rows[3].cells[1].text
            if "（記）" not in cell_text and "(記)" not in cell_text:
                findings.append(Finding(
                    "基本情報", 3,
                    "作成者の名前の後ろに「（記）」が見つからない",
                    "基本情報テーブル: 出席者欄",
                    "議事録作成者の名前の後ろに「（記）」を追加する"
                ))
    return findings


def check_decision_todo_style(doc):
    """Check #4: 決定事項/ToDoの文末が「～する。」or 体言止め."""
    findings = []
    for table in doc.tables:
        if len(table.columns) >= 2:
            header = table.rows[0].cells[1].text.strip() if len(table.rows) > 0 else ""
            if header in ("決定事項", "To Do"):
                for ri, row in enumerate(table.rows[1:], 1):
                    text = row.cells[1].text.strip()
                    if text and not text.endswith("。") and not _is_taigen_dome(text):
                        findings.append(Finding(
                            "決定事項/ToDo", 4,
                            f"{header}の記載が「～する。」または体言止めで終わっていない",
                            f"{header}テーブル 行{ri}: 「{text[:40]}...」",
                            "「～する。」で統一するか、体言止めで統一する"
                        ))
    return findings


def _is_taigen_dome(text):
    """Check if text ends with a noun-like pattern (体言止め).

    Excludes common verb endings (る/す/く/ぐ/む/ぶ/つ/ぬ/う preceded by hiragana)
    and checks for noun-like suffixes commonly used in business minutes.
    """
    # Reject if ends with verb-like pattern (hiragana + verb ending)
    if re.search(r"[いきしちにひみりえけせてねへめれげぜでべ][るすくぐむぶつぬう]$", text):
        return False
    # Accept if ends with common noun suffixes in business context
    return bool(re.search(
        r"(?:[ンめト務事項更変更定認証析]|こと|もの|のみ|予定|完了|対象|方針|結果|内容|状況|概要|報告|確認|承認|調査|検討|対応|作成|実施|送付|提出|共有)$",
        text
    ))


def check_polite_form(paragraphs):
    """Check #10: 常体で統一されているか."""
    findings = []
    for idx, text in paragraphs:
        for pattern, suggestion in POLITE_PATTERNS:
            matches = re.findall(pattern, text)
            if matches:
                findings.append(Finding(
                    "議事", 10,
                    f"敬体表現が使われている: 「{'、'.join(matches)}」",
                    f"段落{idx}: 「{text[:50]}...」",
                    f"「{suggestion}」のように常体に変換する"
                ))
    return findings


def check_tags(doc, paragraphs):
    """Check #11: 決定事項/ToDo タグの対応."""
    findings = []
    # Count tags in paragraphs
    todo_tags = set()
    decision_tags = set()
    for _, text in paragraphs:
        for m in re.finditer(r"<ToDo#(\d+)", text):
            todo_tags.add(int(m.group(1)))
        for m in re.finditer(r"<決定事項#(\d+)", text):
            decision_tags.add(int(m.group(1)))

    # Count table rows
    todo_count = 0
    decision_count = 0
    for table in doc.tables:
        if len(table.columns) >= 2 and len(table.rows) > 0:
            header = table.rows[0].cells[1].text.strip()
            if header == "決定事項":
                decision_count = len(table.rows) - 1
            elif header == "To Do":
                todo_count = len(table.rows) - 1

    if todo_count > 0 and not todo_tags:
        findings.append(Finding(
            "議事", 11,
            f"ToDoテーブルに{todo_count}件のToDoがあるが、議事詳細に<ToDo#N>タグが見つからない",
            "議事詳細",
            "各ToDoに対応する議事詳細の箇所に<ToDo#N （担当者）>タグを追加する"
        ))
    if decision_count > 0 and not decision_tags:
        findings.append(Finding(
            "議事", 11,
            f"決定事項テーブルに{decision_count}件の決定事項があるが、議事詳細に<決定事項#N>タグが見つからない",
            "議事詳細",
            "各決定事項に対応する議事詳細の箇所に<決定事項#N>タグを追加する"
        ))
    return findings


def check_demonstratives(paragraphs):
    """Check #14: 指示代名詞の使用."""
    findings = []
    for idx, text in paragraphs:
        for pattern in DEMONSTRATIVE_PATTERNS:
            matches = re.findall(pattern, text)
            if matches:
                findings.append(Finding(
                    "議事", 14,
                    f"指示代名詞が使われている: 「{'、'.join(set(matches))}」",
                    f"段落{idx}: 「{text[:50]}...」",
                    "具体的な名詞・固有名詞に置き換える"
                ))
                break  # one finding per paragraph
    return findings


def check_colloquial(paragraphs):
    """Check #15: 口語表現の検出."""
    findings = []
    for idx, text in paragraphs:
        for colloquial, replacement in COLLOQUIAL_MAP.items():
            if colloquial in text:
                suggestion = f"「{replacement}」に置き換える" if replacement else "ビジネス表現に言い換える"
                findings.append(Finding(
                    "議事", 15,
                    f"口語/カジュアル表現: 「{colloquial}」",
                    f"段落{idx}: 「{text[:50]}...」",
                    suggestion
                ))
    return findings


def check_vague_time(paragraphs):
    """Check #7: 曖昧な時期表現."""
    findings = []
    for idx, text in paragraphs:
        for pattern in VAGUE_TIME_PATTERNS:
            match = re.search(pattern, text)
            if match:
                findings.append(Finding(
                    "議事", 7,
                    f"時期が不明確: 「{match.group()}」",
                    f"段落{idx}: 「{text[:50]}...」",
                    "具体的な日付（YYYY/MM/DD）または期間を記載する"
                ))
    return findings


def check_sentence_length(paragraphs):
    """Check #18: 1文3行以内（約120文字）."""
    findings = []
    for idx, text in paragraphs:
        sentences = re.split(r"。", text)
        for sentence in sentences:
            if len(sentence) > 120:
                findings.append(Finding(
                    "表現", 18,
                    f"1文が長すぎる（{len(sentence)}文字）",
                    f"段落{idx}: 「{sentence[:60]}...」",
                    "箇条書きに分割するか、接続詞で文を分ける"
                ))
    return findings


def check_rude_expressions(paragraphs):
    """Check #19: 配慮に欠ける表現."""
    findings = []
    for idx, text in paragraphs:
        for pattern in RUDE_PATTERNS:
            match = re.search(pattern, text)
            if match:
                findings.append(Finding(
                    "表現", 19,
                    f"参加者への配慮に欠ける表現: 「{match.group()}」",
                    f"段落{idx}: 「{text[:50]}...」",
                    "表現を改める（例: 「レベルが低い」→「改善すべき点もあるが」）"
                ))
    return findings


def check_about_prefix(paragraphs):
    """Check: 「～について」「～に対して」で文を始めていないか."""
    findings = []
    for idx, text in paragraphs:
        if re.match(r"^.{1,15}(?:について|に対して)[、,]", text):
            findings.append(Finding(
                "表現", 0,
                "「～について」「～に対して」で文を始めている",
                f"段落{idx}: 「{text[:50]}...」",
                "主語＋述語の構造に書き換える（例:「PQ会議について→PQ会議の定義を」）"
            ))
    return findings


def check_missing_subject(paragraphs):
    """Check #13: 主語の省略（議事詳細の発言文で主語が省略されているケースを検出）."""
    findings = []
    for idx, text in paragraphs:
        # Skip section headings, agenda items, speaker-only lines, tags
        if text.startswith("<") or text.startswith("（") or text.startswith("日時") \
                or text.startswith("場所") or text.startswith("出席者") \
                or re.match(r"^[^。]{0,20}$", text):  # very short = heading/label
            continue
        sentences = re.split(r"。", text)
        for s in sentences:
            s = s.strip()
            # Only flag real content sentences (30+ chars, no particles at all)
            if len(s) > 30 and not re.search(r"[はがをにでの]", s):
                findings.append(Finding(
                    "議事", 13,
                    f"主語/助詞が見当たらない文: 「{s[:50]}...」",
                    f"段落{idx}",
                    "5W1Hを意識して主語・目的語を補う"
                ))
    return findings


def check_additional_expressions(paragraphs):
    """Check additional expression rules (～ので、本/当、擬似熟語、および、鑑みる、基本的に、お/ご)."""
    findings = []
    for idx, text in paragraphs:
        for pattern, message, suggestion in ADDITIONAL_EXPRESSION_CHECKS:
            match = re.search(pattern, text)
            if match:
                findings.append(Finding(
                    "表現", 0,
                    f"{message}: 「{match.group()}」",
                    f"段落{idx}: 「{text[:50]}...」",
                    suggestion
                ))
    return findings


def run_review(input_path):
    """Run all checks and return findings."""
    doc = Document(str(input_path))
    paragraphs = extract_text_paragraphs(doc)

    all_findings = []
    all_findings.extend(check_font(doc))
    all_findings.extend(check_recorder_mark(doc))
    all_findings.extend(check_decision_todo_style(doc))
    all_findings.extend(check_polite_form(paragraphs))
    all_findings.extend(check_tags(doc, paragraphs))
    all_findings.extend(check_demonstratives(paragraphs))
    all_findings.extend(check_colloquial(paragraphs))
    all_findings.extend(check_vague_time(paragraphs))
    all_findings.extend(check_sentence_length(paragraphs))
    all_findings.extend(check_rude_expressions(paragraphs))
    all_findings.extend(check_about_prefix(paragraphs))
    all_findings.extend(check_missing_subject(paragraphs))
    all_findings.extend(check_additional_expressions(paragraphs))

    return all_findings


def generate_review_docx(findings, output_path, source_name=""):
    """Generate review result as a Word document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # Title
    title_p = doc.add_paragraph()
    run = title_p.add_run("議事録レビュー結果")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    if source_name:
        p = doc.add_paragraph()
        run = p.add_run(f"対象ファイル: {source_name}")
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    p = doc.add_paragraph()
    run = p.add_run(f"指摘件数: {len(findings)}件")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    doc.add_paragraph()

    if not findings:
        p = doc.add_paragraph()
        run = p.add_run("指摘事項はありません。チェックリスト全項目PASS。")
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    else:
        # Create findings table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header
        headers = ["No", "分類", "指摘事項", "該当箇所", "改善案"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(h)
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.bold = True
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            # Shade header
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "pct20")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "auto")
            tcPr.append(shd)

        # Data rows
        for i, f in enumerate(findings, 1):
            row = table.add_row()
            texts = [str(i), f.category, f.text, f.location, f.suggestion]
            for ci, txt in enumerate(texts):
                cell = row.cells[ci]
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(txt)
                run.font.name = FONT_NAME
                run.font.size = Pt(9)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    doc.save(str(output_path))
    print(f"Review report: {output_path} ({len(findings)} findings)")
    return len(findings)


def main():
    parser = argparse.ArgumentParser(description="Review meeting minutes .docx")
    parser.add_argument("--input", required=True, help="Input .docx to review")
    parser.add_argument("--output", required=True, help="Output review report .docx")
    args = parser.parse_args()

    findings = run_review(args.input)
    source_name = Path(args.input).name
    generate_review_docx(findings, args.output, source_name)


if __name__ == "__main__":
    main()
