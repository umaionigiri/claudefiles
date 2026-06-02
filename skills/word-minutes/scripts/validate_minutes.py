#!/usr/bin/env python3
"""
validate_minutes.py — 議事録 .docx の構造バリデーション

generate_minutes.py で生成した .docx が正しく構成されているか検証する。
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

FONT_NAME = "Meiryo UI"
REQUIRED_SECTIONS = ["<アジェンダ>", "<決定事項>", "<ToDos>", "<議事詳細>"]
OPTIONAL_SECTIONS = ["<次回ミーティング>"]


class Check:
    def __init__(self, name, passed, detail=""):
        self.name = name
        self.passed = passed
        self.detail = detail

    def __str__(self):
        status = "PASS" if self.passed else "FAIL"
        msg = f"  [{status}] {self.name}"
        if self.detail:
            msg += f" — {self.detail}"
        return msg


def validate(docx_path):
    """Run all validation checks. Returns (checks, pass_count, fail_count)."""
    doc = Document(str(docx_path))
    checks = []

    # 1. File size
    file_size = Path(docx_path).stat().st_size
    checks.append(Check("ファイルサイズ", file_size > 5000,
                         f"{file_size} bytes"))

    # 2. Required sections
    all_text = "\n".join(p.text for p in doc.paragraphs)
    for section in REQUIRED_SECTIONS:
        found = section in all_text
        checks.append(Check(f"セクション: {section}", found,
                             "見つかりました" if found else "見つかりません"))

    # 3. Tables exist
    table_count = len(doc.tables)
    checks.append(Check("テーブル数", table_count >= 1,
                         f"{table_count}個のテーブル"))

    # 4. Basic info table structure (first table should be 5x2)
    if table_count >= 1:
        t = doc.tables[0]
        correct = len(t.rows) >= 5 and len(t.columns) >= 2
        checks.append(Check("基本情報テーブル構造", correct,
                             f"{len(t.rows)}行x{len(t.columns)}列"))
    else:
        checks.append(Check("基本情報テーブル構造", False, "テーブルなし"))

    # 5. Font check (>= 80% Meiryo UI)
    total_runs = 0
    meiryo_runs = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text.strip():
                total_runs += 1
                fname = run.font.name
                if fname is None or fname == FONT_NAME:
                    meiryo_runs += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip():
                            total_runs += 1
                            fname = run.font.name
                            if fname is None or fname == FONT_NAME:
                                meiryo_runs += 1

    if total_runs > 0:
        pct = meiryo_runs / total_runs * 100
        checks.append(Check("フォント Meiryo UI", pct >= 80,
                             f"{meiryo_runs}/{total_runs} ({pct:.0f}%)"))
    else:
        checks.append(Check("フォント Meiryo UI", False, "テキストrunなし"))

    # 6. Ends with 「以上」
    last_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    ends_with_ijou = last_texts[-1] == "以上" if last_texts else False
    checks.append(Check("文書末尾「以上」", ends_with_ijou,
                         f"末尾: 「{last_texts[-1][:20]}」" if last_texts else "空文書"))

    # 7. Paragraph count
    non_empty = len([p for p in doc.paragraphs if p.text.strip()])
    checks.append(Check("段落数", non_empty >= 10,
                         f"{non_empty}段落"))

    # 8. 常体チェック (no です/ます in main body)
    polite_count = 0
    for p in doc.paragraphs:
        text = p.text
        if text.startswith("<") and text.endswith(">"):
            continue  # skip section headings
        polite_count += len(re.findall(r"(?:です|ます)[。、\s]", text))
    checks.append(Check("常体統一", polite_count == 0,
                         f"敬体表現: {polite_count}箇所" if polite_count > 0 else ""))

    # 9. ToDo/決定事項 tags in discussion section
    in_discussion = False
    tag_count = 0
    for p in doc.paragraphs:
        if "<議事詳細>" in p.text:
            in_discussion = True
        elif p.text.startswith("<") and p.text.endswith(">") and in_discussion:
            in_discussion = False
        if in_discussion:
            tag_count += len(re.findall(r"<(?:ToDo|決定事項)#\d+", p.text))
    # Only check if there are decisions/todos
    has_content = False
    for table in doc.tables:
        if len(table.rows) > 1 and len(table.columns) >= 2:
            header = table.rows[0].cells[1].text.strip() if table.rows[0].cells else ""
            if header in ("決定事項", "To Do"):
                has_content = True
    if has_content:
        checks.append(Check("ToDo/決定事項タグ", tag_count > 0,
                             f"{tag_count}個のタグ"))

    # Summary
    pass_count = sum(1 for c in checks if c.passed)
    fail_count = sum(1 for c in checks if not c.passed)

    return checks, pass_count, fail_count


def main():
    parser = argparse.ArgumentParser(description="Validate meeting minutes .docx")
    parser.add_argument("docx", help=".docx file to validate")
    args = parser.parse_args()

    checks, pass_count, fail_count = validate(args.docx)

    print(f"\nValidation: {Path(args.docx).name}")
    print("=" * 50)
    for c in checks:
        print(c)
    print("=" * 50)
    print(f"Result: {pass_count} PASS, {fail_count} FAIL")

    sys.exit(0 if fail_count == 0 else 1)


if __name__ == "__main__":
    main()
