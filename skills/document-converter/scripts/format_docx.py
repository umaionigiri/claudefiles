#!/usr/bin/env python3
"""format_docx.py - DOCX formatting script for Claude Code"""
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK


def set_run_font(run, font_name="Meiryo UI", size=10.5):
    """Set font for a run with East Asian support."""
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def set_cell_shading(cell, color="F0F0F0"):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def format_docx(filepath):
    doc = Document(filepath)

    # H2見出しの前に改ページ挿入（H1直後は除く）
    paragraphs = list(doc.paragraphs)
    for i, para in enumerate(paragraphs):
        if para.style.name == 'Heading 2' and i > 0:
            prev_para = paragraphs[i - 1]
            if prev_para.style.name != 'Heading 1':
                run = para.insert_paragraph_before().add_run()
                run.add_break(WD_BREAK.PAGE)

    # 段落のフォント設定
    for para in doc.paragraphs:
        for run in para.runs:
            if para.style.name.startswith('Heading'):
                size = 14 if '1' in para.style.name else 12
                set_run_font(run, "Meiryo UI", size)
            else:
                set_run_font(run, "Meiryo UI", 10.5)

    # 表の罫線設定・フォント適用・タイトル行色付け
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if i == 0:
                    set_cell_shading(cell, "F0F0F0")
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run, "Meiryo UI", 10)

    doc.save(filepath)
    print(f"Formatted: {filepath}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 format_docx.py <file.docx>")
        sys.exit(1)
    format_docx(sys.argv[1])
